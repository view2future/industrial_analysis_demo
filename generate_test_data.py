#!/usr/bin/env python3
"""
Generate test data files with 10,000 POI entries for testing
"""

import json
import pandas as pd
import random
from docx import Document
from docx.shared import Inches
import os
from datetime import datetime

# Create test data directory if it doesn't exist
os.makedirs('test_data', exist_ok=True)

# List of sample companies, universities, and research institutes
companies = [
    "华为技术有限公司", "阿里巴巴集团", "腾讯科技有限公司", "百度公司", "京东集团", 
    "美团点评", "字节跳动", "滴滴出行", "小米科技", "联想集团", "海尔集团", "格力电器",
    "美的集团", "比亚迪", "宁德时代", "隆基绿能", "阳光电源", "汇川技术", "恒瑞医药"
]

universities = [
    "北京大学", "清华大学", "复旦大学", "上海交通大学", "浙江大学", "中国科学技术大学",
    "南京大学", "华中科技大学", "中山大学", "西安交通大学", "北京航空航天大学", 
    "北京理工大学", "哈尔滨工业大学", "西北工业大学", "电子科技大学", "北京邮电大学",
    "大连理工大学", "东南大学", "天津大学", "南开大学"
]

research_institutes = [
    "中国科学院", "中国社会科学院", "中国工程院", "中科院计算技术研究所", "中科院自动化研究所",
    "中科院半导体研究所", "中科院物理研究所", "中科院化学研究所", "中科院生物物理研究所", 
    "中科院上海光机所", "中科院光电技术研究所", "中国医学科学院", "中国农业科学院",
    "中国林业科学研究院", "中国水产科学研究院", "中国热带农业科学院", "中国建筑材料科学研究总院",
    "中国钢研科技集团", "中国有研科技集团", "中国航天科技集团"
]

# Chinese provinces and cities
regions = [
    "北京市", "上海市", "天津市", "重庆市", "河北省", "山西省", "辽宁省", "吉林省", 
    "黑龙江省", "江苏省", "浙江省", "安徽省", "福建省", "江西省", "山东省", "河南省", 
    "湖北省", "湖南省", "广东省", "海南省", "四川省", "贵州省", "云南省", "陕西省", 
    "甘肃省", "青海省", "台湾省", "内蒙古自治区", "广西壮族自治区", "西藏自治区", 
    "宁夏回族自治区", "新疆维吾尔自治区", "香港特别行政区", "澳门特别行政区"
]

def generate_poi_data(count=10000):
    """Generate POI data with specified count."""
    pois = []
    
    for i in range(count):
        # Randomly choose entity type
        entity_type = random.choice(['企业', '高校', '科研院所'])
        
        if entity_type == '企业':
            name = random.choice(companies) + f"有限公司{i+1:04d}"
            description = "领先的科技企业"
        elif entity_type == '高校':
            name = f"{random.choice(universities)}第{i+1:04d}分校"
            description = "高等教育机构"
        else:  # 研究院所
            name = f"{random.choice(research_institutes)}第{i+1:04d}分院"
            description = "科研机构"
        
        # Generate random address
        region = random.choice(regions)
        street_num = random.randint(1, 999)
        address = f"{region}某区某街道{street_num}号"
        
        poi = {
            "name": name,
            "type": entity_type,
            "address": address,
            "region": region,
            "description": description
        }
        
        pois.append(poi)
    
    return pois

# Generate 10,000 POIs
print("Generating 10,000 POI entries...")
pois = generate_poi_data(10000)

# 1. Create JSON file
json_data = {"pois": pois}
with open('test_data/test_data_large.json', 'w', encoding='utf-8') as f:
    json.dump(json_data, f, ensure_ascii=False, indent=2)

print("JSON file created: test_data/test_data_large.json")

# 2. Create Excel file
df = pd.DataFrame(pois)
df.to_excel('test_data/test_data_large.xlsx', index=False)

print("Excel file created: test_data/test_data_large.xlsx")

# 3. Create CSV file (as another format)
df.to_csv('test_data/test_data_large.csv', index=False, encoding='utf-8')

print("CSV file created: test_data/test_data_large.csv")

# 4. Create DOCX file
doc = Document()
doc.add_heading('企业、高校、科研院所清单', 0)

# Add a table with headers
table = doc.add_table(rows=1, cols=5)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '序号'
hdr_cells[1].text = '名称'
hdr_cells[2].text = '类型'
hdr_cells[3].text = '地址'
hdr_cells[4].text = '行政区域'

# Add data rows (first 100 for demo)
for i, poi in enumerate(pois[:100]):  # Using only first 100 to keep the file manageable
    row_cells = table.add_row().cells
    row_cells[0].text = str(i + 1)
    row_cells[1].text = poi['name']
    row_cells[2].text = poi['type']
    row_cells[3].text = poi['address']
    row_cells[4].text = poi['region']

doc.save('test_data/test_data_large.docx')

print("DOCX file created: test_data/test_data_large.docx")

# 5. Create a sample text file
with open('test_data/test_data_large.txt', 'w', encoding='utf-8') as f:
    f.write("企业、高校、科研院所清单\n\n")
    for i, poi in enumerate(pois[:1000]):  # Using first 1000 for text file
        f.write(f"{i+1}. 名称: {poi['name']}\n")
        f.write(f"   类型: {poi['type']}\n")
        f.write(f"   地址: {poi['address']}\n")
        f.write(f"   行政区域: {poi['region']}\n")
        f.write(f"   描述: {poi['description']}\n")
        f.write("\n")

print("TXT file created: test_data/test_data_large.txt")

print(f"\nGenerated {len(pois)} POI entries in multiple formats")
print("Files created in test_data/ directory:")
for file in os.listdir('test_data'):
    path = os.path.join('test_data', file)
    size = os.path.getsize(path)
    print(f"  - {file} ({size:,} bytes)")

print("\nTest data generation completed!")