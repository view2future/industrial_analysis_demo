#!/usr/bin/env python3
"""
Generate realistic test data files for POI visualization
Uses real Chinese enterprises, universities, and research institutes with real addresses
"""

import json
import pandas as pd
import random
import requests
import time
import os
from docx import Document
from docx.shared import Inches

# Create test data directory if it doesn't exist
os.makedirs('test_data_real', exist_ok=True)

# Configuration
BAIDU_AK = "7d56c02f1d2b48a9af5b7d62bb08b62e"  # From config.json
GEOCODE_URL = "http://api.map.baidu.com/geocoding/v3/"

def get_coordinates(address, region=""):
    """Get coordinates for an address using Baidu Map API."""
    try:
        params = {
            'address': address,
            'output': 'json',
            'ak': BAIDU_AK
        }
        if region:
            params['region'] = region

        response = requests.get(GEOCODE_URL, params=params, timeout=10)

        # Check if response is valid JSON
        if response.status_code != 200:
            print(f"HTTP error for '{address}': {response.status_code}")
            return None

        try:
            result = response.json()
        except ValueError:
            # Handle case where response is not JSON
            print(f"Non-JSON response for '{address}': {response.text[:200]}...")
            return None

        # Check for successful response
        if result.get('status') == 0:  # Success
            location = result['result']['location']
            return {
                'longitude': location['lng'],
                'latitude': location['lat']
            }
        else:
            # Handle API error
            message = result.get('message', f"API Error Code: {result.get('status')}")
            print(f"Geocoding failed for '{address}': {message}")
            return None
    except Exception as e:
        print(f"Error geocoding '{address}': {e}")
        return None

# Real data for Chinese companies, universities, and research institutes
real_companies = [
    {"name": "华为技术有限公司", "address": "广东省深圳市龙岗区坂田华为基地"},
    {"name": "腾讯科技有限公司", "address": "广东省深圳市南山区科技园科技中一路腾讯大厦"},
    {"name": "阿里巴巴集团", "address": "浙江省杭州市西湖区文一西路969号"},
    {"name": "百度公司", "address": "北京市海淀区上地十街10号"},
    {"name": "京东集团", "address": "北京市亦庄经济开发区科创十一街18号"},
    {"name": "小米科技有限责任公司", "address": "北京市海淀区清河中街68号"},
    {"name": "美团网", "address": "北京市朝阳区望京东园四区1号"},
    {"name": "字节跳动", "address": "北京市海淀区知春路甲48号"},
    {"name": "滴滴出行", "address": "北京市海淀区东北旺西路8号"},
    {"name": "新浪公司", "address": "北京市海淀区西北旺东路10号院"},
    {"name": "搜狐公司", "address": "北京市海淀区科学院南路2号"},
    {"name": "网易公司", "address": "广东省广州市天河区思蕴路5号"},
    {"name": "携程旅行网", "address": "上海市闵行区联明路555号"},
    {"name": "盛大网络", "address": "上海市浦东新区张江高科技园区郭守敬路351号"},
    {"name": "巨人网络", "address": "上海市松江区中辰路655号"},
    {"name": "用友网络科技股份有限公司", "address": "北京市海淀区北清路68号"},
    {"name": "东软集团", "address": "辽宁省大连市软件园东路21号"},
    {"name": "中科软科技股份有限公司", "address": "北京市海淀区中关村南大街34号"},
    {"name": "广联达科技股份有限公司", "address": "北京市海淀区西北旺东路中关村软件园二期"},
    {"name": "恒生电子股份有限公司", "address": "浙江省杭州市滨江区江南大道3588号"}
]

real_universities = [
    {"name": "北京大学", "address": "北京市海淀区颐和园路5号"},
    {"name": "清华大学", "address": "北京市海淀区清华园1号"},
    {"name": "复旦大学", "address": "上海市杨浦区邯郸路220号"},
    {"name": "上海交通大学", "address": "上海市闵行区东川路800号"},
    {"name": "浙江大学", "address": "浙江省杭州市西湖区余杭塘路866号"},
    {"name": "南京大学", "address": "江苏省南京市鼓楼区汉口路22号"},
    {"name": "中山大学", "address": "广东省广州市新港西路135号"},
    {"name": "华中科技大学", "address": "湖北省武汉市洪山区珞喻路1037号"},
    {"name": "西安交通大学", "address": "陕西省西安市雁翔路99号"},
    {"name": "哈尔滨工业大学", "address": "黑龙江省哈尔滨市南岗区西大直街92号"},
    {"name": "北京航空航天大学", "address": "北京市海淀区学院路37号"},
    {"name": "北京理工大学", "address": "北京市海淀区中关村南大街5号"},
    {"name": "南开大学", "address": "天津市津南区雅观路135号"},
    {"name": "天津大学", "address": "天津市津南区雅观路135号"},
    {"name": "大连理工大学", "address": "辽宁省大连市甘井子区凌工路2号"},
    {"name": "吉林大学", "address": "吉林省长春市前进大街2699号"},
    {"name": "东北大学", "address": "辽宁省沈阳市和平区文化路三巷11号"},
    {"name": "山东大学", "address": "山东省济南市山大南路27号"},
    {"name": "中国海洋大学", "address": "山东省青岛市崂山区松岭路238号"},
    {"name": "厦门大学", "address": "福建省厦门市思明区思明南路422号"}
]

real_research_institutes = [
    {"name": "中国科学院", "address": "北京市三里河路52号"},
    {"name": "中科院计算技术研究所", "address": "北京市海淀区中关村科学院南路6号"},
    {"name": "中科院自动化研究所", "address": "北京市海淀区中关村东路95号"},
    {"name": "中科院半导体研究所", "address": "北京市海淀区清华东路甲35号"},
    {"name": "中科院物理研究所", "address": "北京市海淀区中关村南三街8号"},
    {"name": "中科院化学研究所", "address": "北京市海淀区中关村北一街2号"},
    {"name": "中科院生物物理研究所", "address": "北京市朝阳区大屯路151号"},
    {"name": "中科院上海光机所", "address": "上海市嘉定区清河路390号"},
    {"name": "中科院光电技术研究所", "address": "四川省成都市双流区光电大道1号"},
    {"name": "中国医学科学院", "address": "北京市东城区东单三条9号"},
    {"name": "中国农业科学院", "address": "北京市海淀区中关村南大街12号"},
    {"name": "中国林业科学研究院", "address": "北京市海淀区东升南路2号"},
    {"name": "中国水产科学研究院", "address": "北京市丰台区永定路100号"},
    {"name": "中国热带农业科学院", "address": "海南省海口市学院路4号"},
    {"name": "中国建筑材料科学研究总院", "address": "北京市朝阳区管庄东里1号"},
    {"name": "中国钢研科技集团", "address": "北京市海淀区学院南路76号"},
    {"name": "中国有研科技集团", "address": "北京市海淀区北三环中路43号"},
    {"name": "中国航天科技集团", "address": "北京市海淀区阜成路16号"},
    {"name": "中国电子科技集团公司", "address": "北京市石景山区双园路11号"},
    {"name": "中国船舶重工集团公司", "address": "北京市海淀区昆明湖南路72号"}
]

def generate_realistic_poi_data(n=100):
    """Generate realistic POI data with actual addresses."""
    pois = []
    
    # Calculate how many of each type to include
    companies_count = n // 3
    universities_count = n // 3
    research_count = n - companies_count - universities_count  # Remaining
    
    # Add companies
    for i in range(companies_count):
        company = random.choice(real_companies)
        pois.append({
            "name": company["name"],
            "type": "企业",
            "address": company["address"],
            "region": company["address"].split("省")[0] + "省" if "省" in company["address"] else company["address"].split("市")[0] + "市",
            "description": "知名科技企业"
        })
    
    # Add universities
    for i in range(universities_count):
        university = random.choice(real_universities)
        pois.append({
            "name": university["name"],
            "type": "高校",
            "address": university["address"],
            "region": university["address"].split("省")[0] + "省" if "省" in university["address"] else university["address"].split("市")[0] + "市",
            "description": "知名高等院校"
        })
    
    # Add research institutes
    for i in range(research_count):
        research = random.choice(real_research_institutes)
        pois.append({
            "name": research["name"],
            "type": "科研院所",
            "address": research["address"],
            "region": research["address"].split("省")[0] + "省" if "省" in research["address"] else research["address"].split("市")[0] + "市",
            "description": "重要科研机构"
        })
    
    # Shuffle the list to randomize order
    random.shuffle(pois)
    
    return pois

def add_coordinates_to_pois(pois):
    """Add coordinates to POIs using Baidu Map API."""
    print("Geocoding addresses using Baidu Map API...")
    geocoded_count = 0
    
    for i, poi in enumerate(pois):
        if geocoded_count >= 10:  # Limit API calls to avoid rate limits
            print("Limited to 10 geocoding operations to respect API limits")
            break
            
        coords = get_coordinates(poi["address"], poi.get("region", ""))
        if coords:
            poi["longitude"] = coords["longitude"]
            poi["latitude"] = coords["latitude"]
            geocoded_count += 1
            print(f"Geocoded {i+1}/{len(pois)}: {poi['name']}")
        else:
            # Use mock coordinates if geocoding fails
            poi["longitude"] = random.uniform(73, 135)
            poi["latitude"] = random.uniform(18, 53)
            print(f"Used mock coordinates for {poi['name']}")
        
        # Respect API rate limits
        time.sleep(0.1)
    
    return pois

def create_json_file(pois, filename):
    """Create JSON file."""
    data = {"pois": pois}
    with open(filename, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    print(f"JSON file created: {filename}")

def create_excel_file(pois, filename):
    """Create Excel file."""
    df = pd.DataFrame(pois)
    df.to_excel(filename, index=False)
    print(f"Excel file created: {filename}")

def create_csv_file(pois, filename):
    """Create CSV file."""
    df = pd.DataFrame(pois)
    df.to_csv(filename, index=False, encoding='utf-8')
    print(f"CSV file created: {filename}")

def create_txt_file(pois, filename):
    """Create text file."""
    with open(filename, 'w', encoding='utf-8') as f:
        f.write("企业、高校、科研院所清单\n\n")
        for i, poi in enumerate(pois):
            f.write(f"{i+1}. 名称: {poi['name']}\n")
            f.write(f"   类型: {poi['type']}\n")
            f.write(f"   地址: {poi['address']}\n")
            f.write(f"   行政区域: {poi['region']}\n")
            f.write(f"   描述: {poi['description']}\n")
            if 'longitude' in poi:
                f.write(f"   经度: {poi['longitude']}\n")
                f.write(f"   纬度: {poi['latitude']}\n")
            f.write("\n")
    print(f"TXT file created: {filename}")

def create_docx_file(pois, filename):
    """Create DOCX file."""
    doc = Document()
    doc.add_heading('企业、高校、科研院所清单', 0)
    
    # Add a table with headers
    table = doc.add_table(rows=1, cols=6)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '序号'
    hdr_cells[1].text = '名称'
    hdr_cells[2].text = '类型'
    hdr_cells[3].text = '地址'
    hdr_cells[4].text = '行政区域'
    hdr_cells[5].text = '描述'
    
    # Add data rows
    for i, poi in enumerate(pois):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i + 1)
        row_cells[1].text = poi['name']
        row_cells[2].text = poi['type']
        row_cells[3].text = poi['address']
        row_cells[4].text = poi['region']
        row_cells[5].text = poi['description']
    
    doc.save(filename)
    print(f"DOCX file created: {filename}")

def create_pdf_file(pois, filename):
    """Create PDF file (using reportlab)."""
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib import colors
    
    doc = SimpleDocTemplate(filename, pagesize=A4)
    styles = getSampleStyleSheet()
    
    # Container for the 'Flowable' objects
    elements = []
    
    # Add title
    title = Paragraph("企业、高校、科研院所清单", styles['Title'])
    elements.append(title)
    
    # Create data for table
    data = [
        ['序号', '名称', '类型', '地址', '行政区域', '描述']
    ]
    
    for i, poi in enumerate(pois):
        data.append([
            str(i+1),
            poi['name'],
            poi['type'],
            poi['address'],
            poi['region'],
            poi['description']
        ])
    
    # Create table
    table = Table(data)
    
    # Add style to table
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 8),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    
    elements.append(table)
    
    # Build the PDF
    doc.build(elements)
    print(f"PDF file created: {filename}")

if __name__ == "__main__":
    print("Generating realistic POI data with real addresses...")
    
    # Generate 100 realistic POIs
    pois = generate_realistic_poi_data(100)
    
    # Add coordinates using Baidu Map API (with rate limiting)
    pois_with_coords = add_coordinates_to_pois(pois)
    
    # Create various file formats
    create_json_file(pois_with_coords, 'test_data_real/poi_data_real.json')
    create_excel_file(pois_with_coords, 'test_data_real/poi_data_real.xlsx')
    create_csv_file(pois_with_coords, 'test_data_real/poi_data_real.csv')
    create_txt_file(pois_with_coords, 'test_data_real/poi_data_real.txt')
    create_docx_file(pois_with_coords, 'test_data_real/poi_data_real.docx')
    create_pdf_file(pois_with_coords, 'test_data_real/poi_data_real.pdf')
    
    print(f"\nCreated test files with {len(pois_with_coords)} realistic POIs")
    print("Files created in test_data_real/ directory:")
    for file in os.listdir('test_data_real'):
        path = os.path.join('test_data_real', file)
        size = os.path.getsize(path)
        print(f"  - {file} ({size:,} bytes)")
    
    # Print summary
    companies = sum(1 for p in pois_with_coords if p['type'] == '企业')
    universities = sum(1 for p in pois_with_coords if p['type'] == '高校')
    research = sum(1 for p in pois_with_coords if p['type'] == '科研院所')
    
    print(f"\nData breakdown:")
    print(f"  - 企业: {companies}")
    print(f"  - 高校: {universities}")
    print(f"  - 科研院所: {research}")
    print(f"  - 总计: {len(pois_with_coords)}")
    
    print("\nRealistic test data generation completed!")