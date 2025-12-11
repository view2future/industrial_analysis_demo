#!/usr/bin/env python3
"""
Enhanced Regional Industrial Analysis Dashboard
Includes LLM-powered report generation, user system, and advanced features
"""

import warnings
warnings.filterwarnings("ignore", message=".*pkg_resources is deprecated.*")

import markdown2

import os
import sys
import json
from leadership_scraper import LeadershipScraper
import logging
from pathlib import Path
from datetime import datetime, timezone
import pytz
from flask import Flask, render_template, request, jsonify, redirect, url_for, flash
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user
from flask_sqlalchemy import SQLAlchemy
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from celery.result import AsyncResult

# Add src to path for imports
sys.path.append(os.path.join(os.path.dirname(__file__), 'src'))

from src.analysis.text_processor import TextProcessor
from src.visualization.dashboard_generator import DashboardGenerator
from src.ai.llm_generator import LLMReportGenerator
from src.tasks.report_tasks import generate_llm_report_task
from src.export.report_exporter import ReportExporter
from src.analysis.sentiment_analyzer import SentimentAnalyzer
from src.analysis.entity_extractor import EntityExtractor
# Optional: policy analyzer (not required for basic parsing)
# from src.analysis.policy_analyzer import PolicyAnalyzer
from src.analysis.investment_evaluator import InvestmentEvaluator
from src.visualization.map_visualizer import MapVisualizer
from src.visualization.knowledge_graph_visualizer import KnowledgeGraphVisualizer
from src.visualization.story_generator import StoryGenerator
from src.analysis.trend_analyzer import TrendAnalyzer
from src.analysis.comparison_analyzer import ComparisonAnalyzer
from src.analysis.poi_parser import PoiDocumentParser
from src.analysis.poi_searcher import BaiduPoiSearcher, PoiDataProcessor, PoiExporter
from src.utils.api_error_handler import api_error_handler, handle_api_error
from src.utils.notification_service import notification_service
from src.utils.performance_optimizer import CacheManager
from src.utils.time_utils import utc_to_beijing, format_beijing_time, now_beijing
from scripts.api_notification_routes import register_notification_routes

# Override the default datetime functions to use Beijing time
def beijing_now():
    """Return current time in Beijing timezone for use in models"""
    return now_beijing()
from src.routes.streaming_routes import streaming_bp
from routes.report_generation import report_gen_bp
from routes.report_generation_new import report_generation_bp
from src.data.wechat_scraper import WeChatScraper
from src.analysis.wechat_content_analyzer import WeChatContentAnalyzer
from functools import lru_cache


# Define project root path
APP_ROOT = Path(__file__).parent.resolve()

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialize Flask app
app = Flask(__name__)
app.secret_key = 'regional_industrial_analysis_2024'

# Suppress jieba pkg_resources deprecation warning
import warnings
warnings.filterwarnings("ignore", category=UserWarning, module="jieba._compat")
app.config['UPLOAD_FOLDER'] = 'data/input'
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB max file size
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///industrial_analysis.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

# Initialize extensions
db = SQLAlchemy(app)
login_manager = LoginManager(app)
login_manager.login_view = 'login'

# Initialize leadership scraper
leadership_scraper = LeadershipScraper()

# Allowed file extensions
ALLOWED_EXTENSIONS = {'txt', 'md', 'json', 'doc', 'docx', 'pdf'}


# Utility function to get current user ID (supporting anonymous users)
def get_current_user_id():
    """Get current user ID, or default ID for anonymous users."""
    if hasattr(current_user, 'is_authenticated') and current_user.is_authenticated:
        return current_user.id
    else:
        # Return default user ID for anonymous users
        return 0


# Database Models
class User(UserMixin, db.Model):
    """User model for authentication."""
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), unique=True, nullable=False)
    password_hash = db.Column(db.String(200), nullable=False)
    role = db.Column(db.String(20), default='user')  # admin, analyst, user
    created_at = db.Column(db.DateTime, default=beijing_now)
    
    def set_password(self, password):
        self.password_hash = generate_password_hash(password)
    
    def check_password(self, password):
        return check_password_hash(self.password_hash, password)


class Report(db.Model):
    """Report model to track generated reports."""
    id = db.Column(db.Integer, primary_key=True)
    report_id = db.Column(db.String(100), unique=True, nullable=False)
    title = db.Column(db.String(200))
    city = db.Column(db.String(50))
    industry = db.Column(db.String(100))
    report_type = db.Column(db.String(20))  # 'llm' or 'upload'
    file_path = db.Column(db.String(300))
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'))
    status = db.Column(db.String(20), default='pending')  # pending, processing, completed, failed
    created_at = db.Column(db.DateTime, default=beijing_now)
    completed_at = db.Column(db.DateTime)


class WeChatArticle(db.Model):
    """WeChat article model to store articles from specified WeChat accounts."""
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(500), nullable=False)
    content = db.Column(db.Text)  # 存储完整文章内容
    publish_date = db.Column(db.String(20))  # 发布日期，格式：YYYY-MM-DD
    author = db.Column(db.String(100))  # 作者
    source_account = db.Column(db.String(200), nullable=False)  # 公众号名称
    url = db.Column(db.String(500))  # 文章链接
    summary = db.Column(db.Text)  # 文章摘要
    keywords = db.Column(db.Text)  # 关键词，以逗号分隔
    industry_relevance = db.Column(db.Text)  # 相关产业，以逗号分隔
    content_html = db.Column(db.Text)  # 原始HTML内容
    created_at = db.Column(db.DateTime, default=beijing_now)  # 创建时间
    updated_at = db.Column(db.DateTime, default=beijing_now, onupdate=beijing_now)  # 更新时间

    def to_dict(self):
        """Convert model instance to dictionary for JSON serialization."""
        return {
            'id': self.id,
            'title': self.title,
            'content': self.content,
            'publish_date': self.publish_date,
            'author': self.author,
            'source_account': self.source_account,
            'url': self.url,
            'summary': self.summary,
            'keywords': self.keywords.split(',') if self.keywords else [],
            'industry_relevance': self.industry_relevance.split(',') if self.industry_relevance else [],
            'content_html': self.content_html,
            'created_at': self.created_at.isoformat() if self.created_at else None,
            'updated_at': self.updated_at.isoformat() if self.updated_at else None
        }


class PolicyAnalysis(db.Model):
    """Policy analysis model to store analyzed policy documents from email URLs."""
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(500), nullable=False)  # 政策标题
    original_url = db.Column(db.String(500))  # 原始URL
    source_type = db.Column(db.String(50), default='email')  # 来源类型
    content = db.Column(db.Text)  # 政策内容
    content_summary = db.Column(db.Text)  # 内容摘要
    analysis_result = db.Column(db.JSON)  # 完整分析结果（JSON格式）
    classification_region = db.Column(db.String(100))  # 区域分类
    classification_industry = db.Column(db.String(200))  # 行业分类
    classification_year = db.Column(db.Integer)  # 年度分类
    classification_policy_type = db.Column(db.String(100))  # 政策类型
    applicability_score = db.Column(db.Float, default=0.0)  # 适用性评分
    entities = db.Column(db.JSON)  # 提取的实体（JSON格式）
    knowledge_graph = db.Column(db.JSON)  # 知识图谱数据（JSON格式）
    llm_interpretation = db.Column(db.JSON)  # LLM解读结果（JSON格式）
    scraped_at = db.Column(db.DateTime, default=beijing_now)  # 抓取时间
    analyzed_at = db.Column(db.DateTime, default=beijing_now)  # 分析时间
    created_at = db.Column(db.DateTime, default=beijing_now)  # 创建时间
    updated_at = db.Column(db.DateTime, default=beijing_now, onupdate=beijing_now)  # 更新时间
    status = db.Column(db.String(50), default='completed')  # 状态：pending, processing, completed, failed
    tags = db.Column(db.String(500))  # 标签，以逗号分隔

    def to_dict(self):
        """Convert model instance to dictionary for JSON serialization."""
        return {
            'id': self.id,
            'title': self.title,
            'original_url': self.original_url,
            'source_type': self.source_type,
            'content_summary': self.content_summary,
            'classification': {
                'region': self.classification_region,
                'industry': self.classification_industry,
                'year': self.classification_year,
                'policy_type': self.classification_policy_type
            },
            'applicability_score': self.applicability_score,
            'scraped_at': self.scraped_at.isoformat() if self.scraped_at else None,
            'analyzed_at': self.analyzed_at.isoformat() if self.analyzed_at else None,
            'created_at': self.created_at.isoformat() if self.created_at else None,
            'updated_at': self.updated_at.isoformat() if self.updated_at else None,
            'status': self.status,
            'tags': self.tags.split(',') if self.tags else []
        }


class QwenProcessingResult(db.Model):
    """Store Qwen text processing results and history"""
    id = db.Column(db.Integer, primary_key=True)
    original_text = db.Column(db.Text, nullable=False)  # Original input text
    processed_text = db.Column(db.Text)  # Processed/modified text
    summary = db.Column(db.Text)  # Summary of the content
    key_points = db.Column(db.JSON)  # Key points extracted as JSON array
    sentiment = db.Column(db.String(50))  # Sentiment (positive, negative, neutral)
    sentiment_score = db.Column(db.Float, default=0.0)  # Sentiment score (-1 to 1)
    entities = db.Column(db.JSON)  # Extracted entities as JSON array
    topics = db.Column(db.JSON)  # Identified topics as JSON array
    suggestions = db.Column(db.JSON)  # Suggestions as JSON array
    processing_type = db.Column(db.String(100), default='analysis')  # Type of processing
    metadata_extra = db.Column(db.JSON)  # Additional metadata as JSON (renamed to avoid SQL Alchemy conflict)
    status = db.Column(db.String(50), default='completed')  # Status: pending, processing, completed, failed
    error_message = db.Column(db.Text)  # Error details if processing failed
    created_at = db.Column(db.DateTime, default=beijing_now)
    updated_at = db.Column(db.DateTime, default=beijing_now, onupdate=beijing_now)

    def to_dict(self):
        """Convert model instance to dictionary for JSON serialization"""
        return {
            'id': self.id,
            'original_text': self.original_text,
            'processed_text': self.processed_text,
            'summary': self.summary,
            'key_points': self.key_points or [],
            'sentiment': self.sentiment,
            'sentiment_score': self.sentiment_score,
            'entities': self.entities or [],
            'topics': self.topics or [],
            'suggestions': self.suggestions or [],
            'processing_type': self.processing_type,
            'metadata': self.metadata_extra or {},  # Map to the renamed field
            'status': self.status,
            'error_message': self.error_message,
            'created_at': self.created_at.isoformat() if self.created_at else None,
            'updated_at': self.updated_at.isoformat() if self.updated_at else None
        }


@login_manager.user_loader
def load_user(user_id):
    try:
        return db.session.get(User, int(user_id))
    except Exception:
        return None


def allowed_file(filename):
    """Check if file extension is allowed."""
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def init_db():
    """Initialize database and create default admin user."""
    with app.app_context():
        db.create_all()

        # Create default admin user if not exists
        if not User.query.filter_by(username='admin').first():
            admin = User(username='admin', role='admin')
            admin.set_password('admin')
            db.session.add(admin)
            db.session.commit()
            logger.info("Default admin user created (username: admin, password: admin)")

        # Ensure PolicyAnalysis table exists
        logger.info("Policy analysis database initialized")


# Authentication Routes
@app.route('/login', methods=['GET', 'POST'])
def login():
    """User login page."""
    if current_user.is_authenticated:
        return redirect(url_for('index'))
    
    # 自动尝试使用admin/admin登录
    admin_user = User.query.filter_by(username='admin').first()
    if admin_user and admin_user.check_password('admin'):
        login_user(admin_user)
        flash('自动登录成功！', 'success')
        return redirect(url_for('index'))
    
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        
        user = User.query.filter_by(username=username).first()
        
        if user and user.check_password(password):
            login_user(user)
            flash('登录成功！', 'success')
            next_page = request.args.get('next')
            return redirect(next_page if next_page else url_for('index'))
        else:
            flash('用户名或密码错误', 'error')
    
    return render_template('login.html')


@app.route('/logout')
@login_required
def logout():
    """User logout."""
    logout_user()
    flash('已退出登录', 'info')
    return redirect(url_for('login'))


@app.route('/register', methods=['GET', 'POST'])
def register():
    """User registration page."""
    if current_user.is_authenticated:
        return redirect(url_for('index'))
    
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        confirm_password = request.form.get('confirm_password')
        
        if password != confirm_password:
            flash('两次输入的密码不一致', 'error')
            return render_template('register.html')
        
        if User.query.filter_by(username=username).first():
            flash('用户名已存在', 'error')
            return render_template('register.html')
        
        user = User(username=username, role='user')
        user.set_password(password)
        db.session.add(user)
        db.session.commit()
        
        flash('注册成功！请登录', 'success')
        return redirect(url_for('login'))
    
    return render_template('register.html')


# Main Application Routes

# Helpers
from typing import Optional

def _resolve_report_file_path(report_id: str, report: 'Report') -> Optional[Path]:
    try:
        if not report:
            return None
        if report.report_type == 'llm':
            default_path = APP_ROOT / 'data' / 'output' / 'llm_reports' / f"{report_id}.json"
            if default_path.exists():
                return default_path
        # Fallback to stored path if valid
        if getattr(report, 'file_path', None) and report.file_path not in ('', '.', None):
            p = Path(report.file_path)
            return p if p.exists() else None
        return None
    except Exception:
        return None
@app.route('/')
@login_required
def index():
    """Main dashboard page with API status and notifications."""
    try:
        # Get user's recent reports from database (limit to 4 most recent)
        recent_reports = Report.query.filter_by(user_id=current_user.id)\
            .order_by(Report.created_at.desc())\
            .limit(4)\
            .all()
        
        # Get user notifications
        user_notifications = notification_service.get_user_notifications(str(current_user.id))
        notification_stats = notification_service.get_notification_stats(str(current_user.id))
        
        # Get API error summary
        api_error_summary = api_error_handler.get_error_summary()
        
        return render_template('index_enhanced.html', 
                             recent_reports=recent_reports,
                             user=current_user,
                             notifications=user_notifications,
                             notification_stats=notification_stats,
                             api_error_summary=api_error_summary)
    except Exception as e:
        logger.error(f"Error in index route: {e}")
        # Even if notifications fail, show the main dashboard
        return render_template('index_enhanced.html', 
                             recent_reports=[],
                             user=current_user,
                             notifications=[],
                             notification_stats={'total': 0, 'unread': 0, 'has_recent_errors': False},
                            api_error_summary={'status': 'ok', 'message': 'No recent errors'})


@app.route('/logo_preview')
@login_required
def logo_preview():
    return render_template('logo_preview.html')

@app.route('/streaming-generate-report', methods=['GET', 'POST'])
def streaming_generate_report():
    """Streaming report generation page with real-time content display."""
    if request.method == 'POST':
        city = request.form.get('city')
        industry = request.form.get('industry')
        additional_context = request.form.get('additional_context', '')
        llm_service = request.form.get('llm_service', 'kimi')
        
        if not city or not industry:
            flash('请输入城市和行业名称', 'error')
            return render_template('generate_report.html')
        
        # Render streaming template with form data
        return render_template('streaming_generate.html',
                             city=city,
                             industry=industry,
                             additional_context=additional_context,
                             llm_service=llm_service)
    
    # GET request - handle query parameters or show basic form
    city = request.args.get('city', '上海')
    industry = request.args.get('industry', '人工智能')
    additional_context = request.args.get('additional_context', '')
    llm_service = request.args.get('llm_service', 'kimi')
    
    return render_template('streaming_generate.html',
                         city=city,
                         industry=industry,
                         additional_context=additional_context,
                         llm_service=llm_service)


@app.route('/api/generate-report', methods=['POST'])
def api_generate_report():
    """API endpoint for generating reports (for AJAX calls)."""
    try:
        data = request.get_json()
        city = data.get('city')
        industry = data.get('industry')
        additional_context = data.get('additional_context', '')
        llm_service = data.get('llm_service', 'kimi')

        if not city or not industry:
            return jsonify({'error': '请输入城市和行业名称'}), 400

        # Create report record
        timestamp = now_beijing().strftime('%Y%m%d_%H%M%S')
        temp_report_id = f"llm_report_{timestamp}"

        report = Report(
            report_id=temp_report_id,
            title=f"{city} {industry} 产业分析报告",
            city=city,
            industry=industry,
            report_type='llm',
            user_id=get_current_user_id(),
            status='processing'
        )
        db.session.add(report)
        db.session.commit()

        # Start background task
        task = generate_llm_report_task.delay(
            city=city,
            industry=industry,
            additional_context=additional_context,
            user_id=get_current_user_id(),
            initial_report_id=temp_report_id,
            llm_service=llm_service,
            app_root_path=str(APP_ROOT)
        )

        return jsonify({
            'task_id': task.id,
            'report_id': temp_report_id,
            'status': 'processing'
        })

    except Exception as e:
        logger.error(f"Error in API generate report: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/policy-analysis')
@login_required
def policy_analysis_dashboard():
    """Policy analysis dashboard page."""
    try:
        # Get unique values for filters
        all_policies = PolicyAnalysis.query.all()

        regions = list(set([p.classification_region for p in all_policies if p.classification_region]))
        industries = list(set([p.classification_industry for p in all_policies if p.classification_industry]))
        years = list(set([p.classification_year for p in all_policies if p.classification_year]))

        # Sort years in descending order
        years.sort(reverse=True)

        return render_template('policy_analysis_dashboard.html',
                             regions=regions,
                             industries=industries,
                             years=years)
    except Exception as e:
        logger.error(f"Error loading policy analysis dashboard: {e}")
        flash('加载政策分析工作台失败', 'error')
        return redirect(url_for('index'))


@app.route('/policy-analysis/detail/<int:policy_id>')
@login_required
def policy_detail_page(policy_id):
    """Redirect to the WeChat-style analysis page for enhanced visualization."""
    return redirect(url_for('policy_wechat_analysis_page', policy_id=policy_id))


@app.route('/api/policy-content-analysis/<int:policy_id>')
@login_required
def api_policy_content_analysis(policy_id):
    """API endpoint for content-driven policy analysis."""
    try:
        from src.visualization.content_driven_viz_engine import ContentDrivenVisualizationEngine

        policy = PolicyAnalysis.query.get_or_404(policy_id)

        # Create content-driven visualizations
        viz_engine = ContentDrivenVisualizationEngine()

        # Analyze the policy content
        content = policy.content or policy.content_summary or policy.title
        analysis = viz_engine.analyze_policy_content(content)
        visualization_data = viz_engine.generate_visualization_data(analysis)

        return jsonify({
            'content_analysis': analysis,
            'visualization_data': visualization_data
        })
    except Exception as e:
        logger.error(f"Error getting policy content analysis: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/policy-stats')
@login_required
def api_policy_stats():
    """API endpoint for policy dashboard statistics."""
    try:
        from datetime import datetime, timedelta

        total = PolicyAnalysis.query.count()

        # Count new policies in the last 7 days
        seven_days_ago = now_beijing() - timedelta(days=7)
        new_this_week = PolicyAnalysis.query.filter(
            PolicyAnalysis.created_at >= seven_days_ago
        ).count()

        # Count unique regions and industries
        regions = db.session.query(PolicyAnalysis.classification_region).distinct().count()
        industries = db.session.query(PolicyAnalysis.classification_industry).distinct().count()

        return jsonify({
            'total': total,
            'newThisWeek': new_this_week,
            'regionsCount': regions,
            'industriesCount': industries
        })
    except Exception as e:
        logger.error(f"Error getting policy stats: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/policies')
@login_required
def api_policies():
    """API endpoint for policies list with filtering and pagination."""
    try:
        page = int(request.args.get('page', 1))
        limit = int(request.args.get('limit', 10))
        region = request.args.get('region', '')
        industry = request.args.get('industry', '')
        year = request.args.get('year', '')
        type_ = request.args.get('type', '')
        search = request.args.get('search', '')

        # Build query
        query = PolicyAnalysis.query

        # Apply filters
        if region:
            query = query.filter(PolicyAnalysis.classification_region == region)
        if industry:
            query = query.filter(PolicyAnalysis.classification_industry == industry)
        if year:
            query = query.filter(PolicyAnalysis.classification_year == int(year))
        if type_:
            query = query.filter(PolicyAnalysis.classification_policy_type == type_)
        if search:
            query = query.filter(
                db.or_(
                    PolicyAnalysis.title.contains(search),
                    PolicyAnalysis.content.contains(search),
                    PolicyAnalysis.content_summary.contains(search)
                )
            )

        # Get total count
        total = query.count()

        # Apply pagination
        policies = query.order_by(PolicyAnalysis.created_at.desc()).offset((page - 1) * limit).limit(limit).all()

        return jsonify({
            'policies': [policy.to_dict() for policy in policies],
            'total': total,
            'page': page,
            'limit': limit
        })
    except Exception as e:
        logger.error(f"Error getting policies: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/policy/<int:policy_id>')
@login_required
def api_policy_detail(policy_id):
    """API endpoint for policy detail."""
    try:
        policy = PolicyAnalysis.query.get_or_404(policy_id)
        return jsonify(policy.to_dict())
    except Exception as e:
        logger.error(f"Error getting policy detail: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/check-email-policies', methods=['POST'])
@login_required
def api_check_email_policies():
    """API endpoint to check email for new policy URLs and process them."""
    try:
        from src.utils.email_reader import EmailReader
        from src.analysis.policy_analysis_integrator import PolicyAnalysisIntegrator

        # Initialize email reader
        email_reader = EmailReader()

        # Get new emails with URLs
        emails_with_urls = email_reader.get_unread_emails_with_urls()

        new_policies_count = 0

        if emails_with_urls:
            # Initialize policy integrator
            integrator = PolicyAnalysisIntegrator()

            for email_info in emails_with_urls:
                for url in email_info['urls']:
                    try:
                        # Analyze the policy from URL
                        analysis_result = integrator.analyze_policy_from_url(url)

                        if analysis_result.get('success'):
                            # Extract classification info
                            classification = analysis_result.get('classification', {})

                            # Create PolicyAnalysis record
                            policy_analysis = PolicyAnalysis(
                                title=analysis_result.get('title', 'Unknown Title'),
                                original_url=analysis_result.get('url'),
                                source_type='email',
                                content=analysis_result.get('content', ''),
                                content_summary=analysis_result.get('content_summary', analysis_result.get('title', ''))[:500],  # Better summary
                                analysis_result=analysis_result,
                                classification_region=classification.get('region', [None])[0] if classification.get('region') else None,
                                classification_industry=classification.get('industry', [None])[0] if classification.get('industry') else None,
                                classification_year=classification.get('year'),
                                classification_policy_type=classification.get('policy_type'),
                                applicability_score=analysis_result.get('policy_analysis', {}).get('applicability', {}).get('score', 0) if analysis_result.get('policy_analysis', {}).get('applicability') else 0,
                                entities=analysis_result.get('entities'),
                                knowledge_graph=analysis_result.get('knowledge_graph'),
                                llm_interpretation=analysis_result.get('llm_interpretation'),
                                tags=','.join(classification.get('region', []) + classification.get('industry', [])),
                                status='completed'
                            )

                            db.session.add(policy_analysis)
                            db.session.commit()
                            new_policies_count += 1

                    except Exception as e:
                        logger.error(f"Error processing policy from email {email_info['id']}, URL {url}: {e}")
                        # Create failed record
                        policy_analysis = PolicyAnalysis(
                            title=f"分析失败 - {url}",
                            original_url=url,
                            source_type='email',
                            status='failed',
                            content_summary=f"分析失败: {str(e)}"
                        )
                        db.session.add(policy_analysis)
                        db.session.commit()

        return jsonify({
            'success': True,
            'new_policies': new_policies_count,
            'message': f'检查完成，新增 {new_policies_count} 条政策'
        })
    except Exception as e:
        logger.error(f"Error checking email policies: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


def make_json_serializable(obj):
    """Recursively convert an object to be JSON serializable."""
    if isinstance(obj, dict):
        return {key: make_json_serializable(value) for key, value in obj.items()}
    elif isinstance(obj, list):
        return [make_json_serializable(item) for item in obj]
    elif isinstance(obj, tuple):
        return tuple(make_json_serializable(item) for item in obj)
    elif callable(obj):
        # Skip callable objects
        return None
    elif hasattr(obj, '__dict__'):
        # Convert objects with __dict__ to dict, recursively
        return make_json_serializable(obj.__dict__)
    else:
        # Try to convert to string as fallback for objects that can't be serialized
        try:
            # Test if it's already JSON serializable
            import json
            json.dumps(obj)
            return obj
        except (TypeError, ValueError):
            # If not serializable, convert to string representation
            return str(obj)

@app.route('/policy-analysis/detail/<int:policy_id>/wechat')
@login_required
def policy_wechat_analysis_page(policy_id):
    """Specialized page for analyzing WeChat policy articles with brain map visualization."""
    try:
        policy = PolicyAnalysis.query.get_or_404(policy_id)

        # Get content-driven visualizations
        from src.visualization.content_driven_viz_engine import ContentDrivenVisualizationEngine
        viz_engine = ContentDrivenVisualizationEngine()

        content = policy.content or policy.content_summary or policy.title
        analysis = viz_engine.analyze_policy_content(content)
        visualization_data = viz_engine.generate_visualization_data(analysis)

        # Prepare classification data
        classification = {
            'region': policy.classification_region,
            'industry': policy.classification_industry,
            'year': policy.classification_year,
            'policy_type': policy.classification_policy_type
        }

        return render_template('policy_wechat_analysis.html',
                             policy=policy,
                             content_analysis=analysis,
                             visualization_data=visualization_data,
                             classification=classification)
    except Exception as e:
        logger.error(f"Error loading WeChat policy analysis page: {e}")
        flash('加载政策分析页面失败', 'error')
        return redirect(url_for('policy_analysis_dashboard'))


@app.route('/qwen-processing')
@login_required
def qwen_processing_page():
    """Page for processing text with Qwen and viewing results."""
    try:
        # Get recent processing results for history
        recent_results = QwenProcessingResult.query.order_by(
            QwenProcessingResult.created_at.desc()
        ).limit(10).all()

        return render_template('qwen_processing.html', recent_results=recent_results)
    except Exception as e:
        logger.error(f"Error loading Qwen processing page: {e}")
        flash('加载Qwen处理页面失败', 'error')
        return redirect(url_for('index'))



@app.route('/api/qwen-result/<int:result_id>')
@login_required
def api_qwen_result(result_id):
    """API endpoint to get Qwen processing result."""
    try:
        result = QwenProcessingResult.query.get_or_404(result_id)
        return jsonify({'success': True, 'result': result.to_dict()})
    except Exception as e:
        logger.error(f"Error getting Qwen result: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/qwen-history')
@login_required
def api_qwen_history_first():
    """API endpoint to get Qwen processing history."""
    try:
        page = int(request.args.get('page', 1))
        limit = int(request.args.get('limit', 20))

        results = QwenProcessingResult.query.order_by(
            QwenProcessingResult.created_at.desc()
        ).offset((page - 1) * limit).limit(limit).all()

        total = QwenProcessingResult.query.count()

        return jsonify({
            'success': True,
            'results': [result.to_dict() for result in results],
            'total': total,
            'page': page,
            'limit': limit
        })
    except Exception as e:
        logger.error(f"Error getting Qwen history: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/policy/<int:policy_id>', methods=['DELETE'])
@login_required
def delete_policy(policy_id):
    """API endpoint to delete a policy."""
    try:
        policy = PolicyAnalysis.query.get_or_404(policy_id)

        # Delete the policy
        db.session.delete(policy)
        db.session.commit()

        return jsonify({
            'success': True,
            'message': '政策删除成功'
        })
    except Exception as e:
        logger.error(f"Error deleting policy: {e}")
        db.session.rollback()
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500




@app.route('/api/policy/<int:policy_id>/reclassify', methods=['POST'])
@login_required
def api_reclassify_policy(policy_id):
    """API endpoint to reclassify a policy."""
    try:
        policy = PolicyAnalysis.query.get_or_404(policy_id)

        # Process with Qwen to get new classification
        from src.ai.qwen_text_processor import QwenCodeTextProcessor
        processor = QwenCodeTextProcessor()

        content = policy.content or policy.content_summary or policy.title
        result = processor.process_text(content, 'policy_classification')

        if result.get('success', True):  # Using success as True for this mock since the processor handles it differently
            # Update classification
            classification = result.get('classification', {})

            policy.classification_region = classification.get('region', [None])[0] if classification.get('region') else None
            policy.classification_industry = classification.get('industry', [None])[0] if classification.get('industry') else None
            policy.classification_year = classification.get('year')
            policy.classification_policy_type = classification.get('policy_type')

            db.session.commit()

            return jsonify({
                'success': True,
                'message': '政策重新分类成功'
            })
        else:
            return jsonify({
                'success': False,
                'error': result.get('error', 'Reclassification failed')
            })

    except Exception as e:
        logger.error(f"Error reclassifying policy: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/api/export-policies', methods=['POST'])
@login_required
def api_export_policies():
    """API endpoint to export selected policies."""
    try:
        import json

        data = request.get_json()
        policy_ids = data.get('policy_ids', [])

        if not policy_ids:
            return jsonify({'success': False, 'error': 'No policy IDs provided'}), 400

        # Get selected policies
        policies = PolicyAnalysis.query.filter(PolicyAnalysis.id.in_(policy_ids)).all()

        # Create export data
        export_data = []
        for policy in policies:
            export_data.append({
                'id': policy.id,
                'title': policy.title,
                'original_url': policy.original_url,
                'content_summary': policy.content_summary,
                'classification': {
                    'region': policy.classification_region,
                    'industry': policy.classification_industry,
                    'year': policy.classification_year,
                    'policy_type': policy.classification_policy_type
                },
                'applicability_score': policy.applicability_score,
                'status': policy.status,
                'created_at': policy.created_at.isoformat() if policy.created_at else None,
                'analyzed_at': policy.analyzed_at.isoformat() if policy.analyzed_at else None
            })

        # Create response with proper headers for download
        from flask import Response
        json_data = json.dumps(export_data, ensure_ascii=False, indent=2)
        return Response(
            json_data,
            mimetype='application/json',
            headers={'Content-Disposition': f'attachment;filename=policies_export_{datetime.now().strftime("%Y%m%d_%H%M%S")}.json'}
        )

    except Exception as e:
        logger.error(f"Error exporting policies: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/api/policy-visualizations')
@login_required
def api_policy_visualizations():
    """API endpoint for policy visualizations data."""
    try:
        # Get region distribution
        region_data = db.session.query(
            PolicyAnalysis.classification_region,
            db.func.count(PolicyAnalysis.id)
        ).filter(
            PolicyAnalysis.classification_region.isnot(None)
        ).group_by(PolicyAnalysis.classification_region).all()

        region_chart = {
            'labels': [r[0] for r in region_data],
            'values': [r[1] for r in region_data]
        }

        # Get industry distribution
        industry_data = db.session.query(
            PolicyAnalysis.classification_industry,
            db.func.count(PolicyAnalysis.id)
        ).filter(
            PolicyAnalysis.classification_industry.isnot(None)
        ).group_by(PolicyAnalysis.classification_industry).all()

        industry_chart = {
            'labels': [i[0] for i in industry_data],
            'values': [i[1] for i in industry_data]
        }

        # Get year distribution
        year_data = db.session.query(
            PolicyAnalysis.classification_year,
            db.func.count(PolicyAnalysis.id)
        ).filter(
            PolicyAnalysis.classification_year.isnot(None)
        ).group_by(PolicyAnalysis.classification_year).order_by(PolicyAnalysis.classification_year).all()

        year_chart = {
            'labels': [str(y[0]) for y in year_data],
            'values': [y[1] for y in year_data]
        }

        # Get policy type distribution
        type_data = db.session.query(
            PolicyAnalysis.classification_policy_type,
            db.func.count(PolicyAnalysis.id)
        ).filter(
            PolicyAnalysis.classification_policy_type.isnot(None)
        ).group_by(PolicyAnalysis.classification_policy_type).all()

        type_chart = {
            'labels': [t[0] for t in type_data],
            'values': [t[1] for t in type_data]
        }

        return jsonify({
            'region_chart': region_chart,
            'industry_chart': industry_chart,
            'year_chart': year_chart,
            'type_chart': type_chart
        })
    except Exception as e:
        logger.error(f"Error getting visualization data: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/generate-report', methods=['GET', 'POST'])
def generate_report():
    """LLM-powered report generation page."""
    if request.method == 'POST':
        city = request.form.get('city')
        industry = request.form.get('industry')
        additional_context = request.form.get('additional_context', '')
        llm_service = request.form.get('llm_service', 'kimi')
        
        if not city or not industry:
            flash('请输入城市和行业名称', 'error')
            return render_template('generate_report.html')
        
        try:
            import uuid as _uuid
            task_id = str(_uuid.uuid4())
            # Redirect to streaming report page; generation will start via SSE
            return redirect(url_for('stream_report_page',
                                  task_id=task_id,
                                  city=city,
                                  industry=industry,
                                  llm_service=llm_service))
            
        except Exception as e:
            logger.error(f"Error starting report generation: {e}")
            flash(f'启动报告生成失败: {str(e)}', 'error')
    
    return render_template('generate_report.html')


@app.route('/task/<task_id>')
@login_required
def task_status_page(task_id):
    """Display task status page."""
    report_id = request.args.get('report_id')
    return render_template('task_status.html', task_id=task_id, report_id=report_id)


@app.route('/task/<task_id>/status')
@app.route('/api/task-status/<task_id>')
def task_status(task_id):
    """Check status of background task (API endpoint)."""
    report_id = request.args.get('report_id')
    
    try:
        from src.tasks.celery_app import celery_app
        task = AsyncResult(task_id, app=celery_app)
        state = task.state
        
        logger.info(f"Task {task_id[:8]}... state from Celery: {state}")
        
        if state == 'PENDING':
            response = {
                'state': state,
                'status': '等待处理...'
            }
        elif state == 'PROGRESS':
            response = {
                'state': state,
                'current': task.info.get('current', 0),
                'total': task.info.get('total', 100),
                'status': task.info.get('status', ''),
                'stage': task.info.get('stage', 'init'),
                'message': task.info.get('message', task.info.get('status', ''))
            }
            logger.info(f"Task PROGRESS - stage: {response['stage']}, message: {response['message']}")
        elif state == 'SUCCESS':
            # Get task result (prefer result over info on SUCCESS)
            task_result = task.result
            logger.info(f"Raw task result: {task_result}")
            logger.info(f"Task result type: {type(task_result)}")
            
            # Check if this is actually a failure disguised as success
            if isinstance(task_result, dict) and not task_result.get('success', True):
                # This is actually a failure - convert to FAILURE state
                logger.info("Task reported SUCCESS but contains error - treating as FAILURE")
                error_msg = task_result.get('error', 'Unknown error')
                response = {
                    'state': 'FAILURE',
                    'status': f'任务失败: {error_msg}',
                    'error': error_msg
                }
                return response
            
            if not isinstance(task_result, dict) or 'report_id' not in task_result:
                # Fallback to info meta if available
                if isinstance(task.info, dict):
                    task_result = task.info
                    logger.info(f"Using task.info as fallback: {task_result}")
                    logger.info(f"Task.info type: {type(task.info)}")
            
            # Extract report_id from result
            result_report_id = None
            if isinstance(task_result, dict):
                result_report_id = task_result.get('report_id')
                logger.info(f"Extracted report_id from task_result: {result_report_id}")
            else:
                logger.info(f"Task result is not a dict or missing report_id: {task_result}")
            
            # Use result report_id or fallback to query parameter
            final_report_id = result_report_id or report_id
            
            logger.info(f"Task SUCCESS - result_report_id: {result_report_id}, query_report_id: {report_id}, final: {final_report_id}")
            
            # Update report status in database
            if final_report_id:
                report = Report.query.filter_by(report_id=final_report_id).first()
                if report:
                    report.status = 'completed'
                    report.completed_at = beijing_now()
                    if isinstance(task_result, dict) and task_result.get('file_path'):
                        report.file_path = task_result.get('file_path')
                    db.session.commit()
                    logger.info(f"Updated report {final_report_id} status to completed")
            
            response = {
                'state': state,
                'status': '完成',
                'result': {
                    'report_id': final_report_id,
                    **(task_result if isinstance(task_result, dict) else {})
                }
            }
        else:
            # Error occurred
            status = str(task.info)
            error_details = None
            
            if isinstance(task.info, dict):
                if 'error' in task.info:
                    status = task.info['error']
                # Check if this is an API error with detailed information
                if 'exc_type' in task.info and 'exc_message' in task.info:
                    error_details = {
                        'error_type': task.info.get('exc_type'),
                        'error_message': task.info.get('exc_message'),
                        'is_api_error': 'api' in task.info.get('exc_type', '').lower()
                    }
            
            response = {
                'state': state,
                'status': status,
                'error_details': error_details
            }
            
            # Log the error for monitoring
            if error_details:
                logger.error(f"Task {task_id[:8]}... failed with API error: {error_details}")
            else:
                logger.error(f"Task {task_id[:8]}... failed: {status}")
        
        return jsonify(response)
        
    except Exception as e:
        logger.error(f"Error checking task status: {e}")
        return jsonify({
            'state': 'ERROR',
            'status': f'检查任务状态失败: {str(e)}'
        })


@app.route('/api/stop-task/<task_id>', methods=['POST'])
@login_required
def stop_task(task_id):
    try:
        from src.tasks.celery_app import celery_app
        celery_app.control.revoke(task_id, terminate=True)
        logger.info(f"Revoked task {task_id}")
        return jsonify({'success': True, 'message': '任务已停止'})
    except Exception as e:
        logger.error(f"Error stopping task {task_id}: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/report/<report_id>')
@login_required
def view_report(report_id):
    """View a generated report."""
    try:
        # Check if user has access to this report
        report = Report.query.filter_by(report_id=report_id).first()
        
        if not report:
            flash('报告不存在', 'error')
            return redirect(url_for('index'))
        
        if report.user_id != current_user.id and current_user.role != 'admin':
            flash('没有权限访问此报告', 'error')
            return redirect(url_for('index'))
        
        # Load report data
        if report.report_type == 'llm':
            # Default expected path
            expected_path = APP_ROOT / 'data' / 'output' / 'llm_reports' / f"{report_id}.json"
            file_path = expected_path
            template = 'report_view_llm.html'
            # Fallback to stored file_path if default missing
            if not file_path.exists() and report.file_path:
                alt_path = Path(report.file_path)
                if alt_path.exists():
                    file_path = alt_path
        else:
            file_path = Path(report.file_path)
            template = 'report_view_upload.html'
        
        if not file_path.exists():
            flash('报告文件不存在', 'error')
            logger.error(f"Report file not found: {file_path}")
            return redirect(url_for('index'))
        
        logger.info(f"Loading report from {file_path}")
        with open(file_path, 'r', encoding='utf-8') as f:
            report_data = json.load(f)
        
        # Read original file content for upload reports
        original_content = None
        if report.report_type == 'upload' and report.file_path:
            try:
                # Extract the original file path from the JSON file path
                # JSON file path is like: data/output/20251105_113016_analysis.json
                # Original file path is like: data/input/20251105_113016_original_filename.txt
                import re
                import os
                json_filename = os.path.basename(report.file_path)
                timestamp_match = re.search(r'(\d{8}_\d{6})', json_filename)
                if timestamp_match:
                    timestamp = timestamp_match.group(1)
                    # Try to find the original file
                    input_dir = os.path.join('data', 'input')
                    if os.path.exists(input_dir):
                        for original_file in os.listdir(input_dir):
                            if original_file.startswith(timestamp):
                                original_file_path = os.path.join(input_dir, original_file)
                                with open(original_file_path, 'r', encoding='utf-8') as f:
                                    original_content = f.read()
                                break
            except Exception as e:
                logger.warning(f"Could not read original content: {e}")

        # For LLM reports, check if the enhanced template should be used
        use_enhanced_template = False
        if report.report_type == 'llm':
            # Check if the report data has enhanced structure (from streaming)
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    temp_data = json.load(f)
                # If the report has the richer structure we created (charts, analysis, etc.), use enhanced template
                if 'charts' in temp_data or 'analysis' in temp_data or 'key_insights' in temp_data:
                    use_enhanced_template = True
            except Exception as e:
                logger.warning(f"Could not check report structure: {e}")
        
        # Determine which template to use
        if use_enhanced_template:
            template = 'report_view_llm_enhanced.html'

        logger.info(f"Rendering report {report_id} with template {template}")
        logger.info(f"Report data keys: {list(report_data.keys())}")
        logger.info(f"Using enhanced template: {use_enhanced_template}")

        # Pre-serialize chart data for templates to avoid tojson issues
        if 'charts' in report_data:
            for chart_name, chart_data in report_data['charts'].items():
                if isinstance(chart_data, dict) and 'data' in chart_data:
                    # Pre-convert to JSON strings for safe template use
                    report_data[f'{chart_name}_json'] = json.dumps(chart_data)
        
        # Add original content to report object for template access
        if original_content:
            report.original_content = original_content
            
        # Add markdown rendering function
        def render_markdown(content):
            if content:
                return markdown2.markdown(
                    content,
                    extras=[
                        "fenced-code-blocks",
                        "tables",
                        "strike",
                        "task_list",
                        "code-friendly",
                        "toc",
                    ],
                )
            return content

        # Convert UTC times to Beijing time for display
        def convert_report_times(obj):
            """Recursively convert datetime strings to Beijing time in report data"""
            if isinstance(obj, dict):
                new_dict = {}
                for key, value in obj.items():
                    if key in ['created_at', 'updated_at', 'completed_at', 'generated_at', 'published_date', 'start_date', 'end_date', 'created_time', 'modified_time', 'timestamp', 'date_created', 'date_updated', 'published_at', 'modified_at']:
                        # Convert UTC datetime to Beijing time
                        if isinstance(value, (int, float)):  # Unix timestamp
                            new_dict[key] = format_beijing_time(value, '%Y-%m-%d %H:%M:%S')
                        elif isinstance(value, str):
                            # Try to parse the date string, assuming it might be in ISO format
                            try:
                                # Handle various date formats
                                import re
                                # Look for date patterns like 2024-11-28T12:49:45 or 2024-11-28 12:49:45
                                date_pattern = r'(\d{4}-\d{2}-\d{2}[T\s]\d{2}:\d{2}:\d{2}(?:\.\d+)?)'
                                match = re.search(date_pattern, value)
                                if match:
                                    dt_str = match.group(1)
                                    if 'T' in dt_str:
                                        date_part, time_part = dt_str.split('T')
                                    else:
                                        date_part, time_part = dt_str.split(' ')

                                    # Split time into components (may have milliseconds)
                                    time_components = time_part.split('.')
                                    base_time = time_components[0]  # Get base time without milliseconds

                                    # Parse the datetime from the extracted string
                                    from datetime import datetime
                                    parsed_dt = datetime.strptime(f"{date_part} {base_time}", '%Y-%m-%d %H:%M:%S')
                                    # Assume this is UTC, convert to Beijing time
                                    import pytz
                                    utc_dt = pytz.UTC.localize(parsed_dt)
                                    beijing_dt = utc_dt.astimezone(pytz.timezone('Asia/Shanghai'))
                                    new_dict[key] = beijing_dt.strftime('%Y-%m-%d %H:%M:%S')
                                else:
                                    # If it doesn't match a recognized date pattern, keep as-is
                                    new_dict[key] = value
                            except Exception:
                                # If parsing fails, keep original value
                                new_dict[key] = value
                        elif hasattr(value, 'strftime'):  # datetime object
                            # This is likely already handled by the DB layer, but just in case
                            if hasattr(value, 'astimezone'):
                                # Convert to Beijing time if timezone-aware
                                beijing_time = format_beijing_time(value, '%Y-%m-%d %H:%M:%S')
                                new_dict[key] = beijing_time
                            else:
                                # Naive datetime - treat as UTC and convert to Beijing time
                                import pytz
                                if value.tzinfo is None:
                                    # If datetime is naive, assume it's UTC
                                    utc_dt = pytz.UTC.localize(value)
                                else:
                                    # If it has timezone info, convert to UTC first then to Beijing
                                    utc_dt = value.astimezone(pytz.UTC)
                                beijing_dt = utc_dt.astimezone(pytz.timezone('Asia/Shanghai'))
                                new_dict[key] = beijing_dt.strftime('%Y-%m-%d %H:%M:%S')
                        else:
                            new_dict[key] = value  # Don't convert if not a time field
                    else:
                        new_dict[key] = convert_report_times(value)  # Recursively process nested structures
                return new_dict
            elif isinstance(obj, list):
                return [convert_report_times(item) for item in obj]
            else:
                return obj

        # Convert report and report_data times to Beijing time
        converted_report_data = convert_report_times(report_data)

        # Convert the report object attributes to Beijing time
        if hasattr(report, 'created_at') and report.created_at:
            try:
                report.created_at_beijing = format_beijing_time(report.created_at, '%Y-%m-%d %H:%M:%S')
            except:
                report.created_at_beijing = str(report.created_at) if report.created_at else 'N/A'

        if hasattr(report, 'completed_at') and report.completed_at:
            try:
                report.completed_at_beijing = format_beijing_time(report.completed_at, '%Y-%m-%d %H:%M:%S')
            except:
                report.completed_at_beijing = str(report.completed_at) if report.completed_at else 'N/A'

        # Add the function to the template context
        return render_template(template,
                             report=report,
                             report_data=converted_report_data,
                             render_markdown=render_markdown)
        
    except json.JSONDecodeError as e:
        logger.error(f"JSON decode error: {e}")
        flash('报告文件格式错误', 'error')
        return redirect(url_for('index'))
    except Exception as e:
        logger.error(f"Error viewing report: {e}")
        import traceback
        logger.error(traceback.format_exc())
        flash('无法加载报告', 'error')
        return redirect(url_for('index'))


@app.route('/upload', methods=['GET', 'POST'])
@login_required
def upload_file():
    """Handle file upload and processing."""
    if request.method == 'POST':
        if 'file' not in request.files:
            flash('没有选择文件', 'error')
            return redirect(request.url)
        
        file = request.files['file']
        if file.filename == '':
            flash('没有选择文件', 'error')
            return redirect(request.url)
        
        if file and allowed_file(file.filename):
            try:
                # Save uploaded file
                filename = secure_filename(file.filename)
                timestamp = now_beijing().strftime('%Y%m%d_%H%M%S')
                filename = f"{timestamp}_{filename}"
                file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                
                # Ensure upload directory exists
                os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
                file.save(file_path)
                
                # Process the file
                processor = TextProcessor()
                analysis_result = processor.analyze_file(file_path)
                
                if analysis_result:
                    # Generate dashboard data
                    dashboard_gen = DashboardGenerator()
                    dashboard_data = dashboard_gen.generate_dashboard_data(analysis_result)
                    
                    # Save processed data
                    output_path = os.path.join('data/output', f"{timestamp}_analysis.json")
                    os.makedirs('data/output', exist_ok=True)
                    
                    # Use custom JSON encoder to handle Plotly figures
                    from plotly.utils import PlotlyJSONEncoder
                    with open(output_path, 'w', encoding='utf-8') as f:
                        json.dump(dashboard_data, f, ensure_ascii=False, indent=2, cls=PlotlyJSONEncoder)
                    
                    logger.info(f"Dashboard data saved to {output_path}")
                    
                    # Create report record
                    report = Report(
                        report_id=f"upload_{timestamp}",
                        title=f"上传文档分析 - {file.filename}",
                        report_type='upload',
                        file_path=output_path,
                        user_id=get_current_user_id(),
                        status='completed',
                        completed_at=beijing_now()
                    )
                    db.session.add(report)
                    db.session.commit()
                    
                    flash('文件上传并分析成功！', 'success')
                    return redirect(url_for('view_report', report_id=report.report_id))
                else:
                    flash('文件处理失败，请检查文件格式', 'error')
                    
            except Exception as e:
                logger.error(f"Error processing file: {e}")
                flash(f'处理文件时出错: {str(e)}', 'error')
        else:
            flash('不支持的文件格式。请上传 .txt, .md, .json, .docx 或 .pdf 文件', 'error')
    
    return render_template('upload.html')


@app.route('/api/reports')
@login_required
def api_reports():
    """API endpoint to get user's reports."""
    reports = Report.query.filter_by(user_id=current_user.id)\
        .order_by(Report.created_at.desc())\
        .all()
    
    return jsonify([{
        'report_id': r.report_id,
        'title': r.title,
        'city': r.city,
        'industry': r.industry,
        'report_type': r.report_type,
        'status': r.status,
        'created_at': r.created_at.isoformat(),
        'completed_at': r.completed_at.isoformat() if r.completed_at else None
    } for r in reports])


@app.route('/api/llm/stream', methods=['POST'])
def api_llm_stream():
    """Stream LLM main content to client in realtime (chunked)."""
    try:
        data = request.get_json()
        city = data.get('city')
        industry = data.get('industry')
        llm_service = data.get('llm_service', 'kimi')
        additional_context = data.get('additional_context', '')
        
        def generate():
            try:
                generator = LLMReportGenerator(llm_service=llm_service)
                for chunk in generator.stream_report_content(city, industry, additional_context):
                    yield chunk
            except Exception as e:
                yield f"\n[错误] {str(e)}"
        
        return app.response_class(generate(), mimetype='text/plain; charset=utf-8')
    except Exception as e:
        logger.error(f"Error starting LLM stream: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/report/<report_id>/qa', methods=['POST'])
@login_required
def api_report_qa(report_id):
    """API endpoint for Q&A about a report."""
    try:
        # Load report
        report = Report.query.filter_by(report_id=report_id).first()
        if not report or (report.user_id != current_user.id and current_user.role != 'admin'):
            return jsonify({'error': '无权限访问'}), 403
        
        question = request.json.get('question')
        if not question:
            return jsonify({'error': '请提供问题'}), 400
        
        # Load report content
        if report.report_type == 'llm':
            file_path = APP_ROOT / 'data' / 'output' / 'llm_reports' / f"{report_id}.json"
        else:
            file_path = Path(report.file_path)
        
        with open(file_path, 'r', encoding='utf-8') as f:
            report_data = json.load(f)
        
        # Get answer from LLM
        generator = LLMReportGenerator()
        answer = generator.answer_question(
            report_data.get('full_content', ''),
            question
        )
        
        return jsonify({'answer': answer})
    
    except Exception as e:
        logger.error(f"Error in Q&A: {e}")
        return jsonify({'error': str(e)}), 500


# API Status and Notification Routes
@app.route('/api/task/<task_id>/status')
def api_task_status(task_id):
    """Get the status of a background task."""
    from src.tasks.celery_app import celery_app
    
    try:
        # Get task result
        task = celery_app.AsyncResult(task_id)
        
        if task.state == 'PENDING':
            # Task is waiting to be processed
            response = {
                'state': task.state,
                'status': '任务正在等待处理...'
            }
        elif task.state == 'PROGRESS':
            # Task is currently being processed
            response = {
                'state': task.state,
                'current': task.info.get('current', 0),
                'total': task.info.get('total', 1),
                'status': task.info.get('status', ''),
                'stage': task.info.get('stage', ''),
                'message': task.info.get('message', '')
            }
        elif task.state == 'SUCCESS':
            # Task completed successfully
            response = {
                'state': task.state,
                'result': task.info
            }
        elif task.state == 'FAILURE':
            # Task failed
            response = {
                'state': task.state,
                'error': str(task.info)
            }
        else:
            # Something else happened
            response = {
                'state': task.state,
                'result': task.info
            }
        
        return jsonify(response)
    
    except Exception as e:
        logger.error(f"Error getting task status: {e}")
        return jsonify({'state': 'FAILURE', 'error': str(e)}), 500


@app.route('/api/report/<report_id>/export/<format>')
@login_required
def export_report(report_id, format):
    """Export report to specified format."""
    try:
        # Load report
        report = Report.query.filter_by(report_id=report_id).first()
        if not report or (report.user_id != current_user.id and current_user.role != 'admin'):
            flash('无权限访问此报告', 'error')
            return redirect(url_for('index'))
        
        # Load report data
        if report.report_type == 'llm':
            file_path = APP_ROOT / 'data' / 'output' / 'llm_reports' / f"{report_id}.json"
        else:
            file_path = Path(report.file_path)
        
        with open(file_path, 'r', encoding='utf-8') as f:
            report_data = json.load(f)
        
        # Export
        exporter = ReportExporter()
        
        if format == 'pdf':
            export_path = exporter.export_to_pdf(report_data, report_id)
            flash(f'PDF报告已生成: {export_path}', 'success')
        elif format == 'word':
            export_path = exporter.export_to_word(report_data, report_id)
            flash(f'Word报告已生成: {export_path}', 'success')
        elif format == 'excel':
            export_path = exporter.export_to_excel(report_data, report_id)
            flash(f'Excel报告已生成: {export_path}', 'success')
        else:
            flash('不支持的导出格式', 'error')
            return redirect(url_for('view_report', report_id=report_id))
        
        return redirect(url_for('view_report', report_id=report_id))
    
    except Exception as e:
        logger.error(f"Error exporting report: {e}")
        flash(f'导出失败: {str(e)}', 'error')
        return redirect(url_for('view_report', report_id=report_id))


@app.route('/api/report/<report_id>/sentiment')
@login_required
def api_sentiment_analysis(report_id):
    """API endpoint for sentiment analysis."""
    try:
        # Load report
        report = Report.query.filter_by(report_id=report_id).first()
        if not report or (report.user_id != current_user.id and current_user.role != 'admin'):
            return jsonify({'error': '无权限访问'}), 403
        
        # Load report content
        if report.report_type == 'llm':
            file_path = APP_ROOT / 'data' / 'output' / 'llm_reports' / f"{report_id}.json"
        else:
            file_path = Path(report.file_path)
        
        with open(file_path, 'r', encoding='utf-8') as f:
            report_data = json.load(f)
        
        # Perform sentiment analysis
        analyzer = SentimentAnalyzer()
        sentiment_result = analyzer.analyze_by_category(report_data)
        risks = analyzer.detect_risks(report_data.get('full_content', ''))
        
        return jsonify({
            'sentiment': sentiment_result,
            'risks': risks
        })
    
    except Exception as e:
        logger.error(f"Error in sentiment analysis: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/report/<report_id>/entities')
@login_required
def api_entity_extraction(report_id):
    """API endpoint for entity extraction."""
    try:
        report = Report.query.filter_by(report_id=report_id).first()
        if not report or (report.user_id != current_user.id and current_user.role != 'admin'):
            return jsonify({'error': '无权限访问'}), 403
        
        # Load report
        if report.report_type == 'llm':
            file_path = APP_ROOT / 'data' / 'output' / 'llm_reports' / f"{report_id}.json"
        else:
            file_path = Path(report.file_path)
        
        with open(file_path, 'r', encoding='utf-8') as f:
            report_data = json.load(f)
        
        # Extract entities
        extractor = EntityExtractor()
        entities = extractor.extract_entities(report_data.get('full_content', ''))
        graph = extractor.build_entity_graph(entities)
        
        return jsonify({
            'entities': entities,
            'graph': graph
        })
    
    except Exception as e:
        logger.error(f"Error in entity extraction: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/report/<report_id>/investment')
@login_required
def api_investment_evaluation(report_id):
    """API endpoint for investment evaluation."""
    try:
        report = Report.query.filter_by(report_id=report_id).first()
        if not report or (report.user_id != current_user.id and current_user.role != 'admin'):
            return jsonify({'error': '无权限访问'}), 403
        
        # Load report
        if report.report_type == 'llm':
            file_path = APP_ROOT / 'data' / 'output' / 'llm_reports' / f"{report_id}.json"
        else:
            file_path = Path(report.file_path)
        
        with open(file_path, 'r', encoding='utf-8') as f:
            report_data = json.load(f)
        
        # Evaluate investment
        evaluator = InvestmentEvaluator()
        evaluation = evaluator.evaluate(report_data)
        
        return jsonify(evaluation)
    
    except Exception as e:
        logger.error(f"Error in investment evaluation: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/report/<report_id>/visualizations')
@login_required
def api_report_visualizations(report_id):
    """API endpoint for generating map visualizations."""
    try:
        report = Report.query.filter_by(report_id=report_id).first()
        if not report or (report.user_id != current_user.id and current_user.role != 'admin'):
            return jsonify({'error': '无权限访问'}), 403
        
        # Load report (robust path handling)
        file_path = _resolve_report_file_path(report_id, report)
        if not file_path:
            return jsonify({'error': '报告文件不存在'}), 404
        with open(file_path, 'r', encoding='utf-8') as f:
            report_data = json.load(f)
        
        # Load config
        with open('config.json', 'r', encoding='utf-8') as f:
            config = json.load(f)
        
        # Generate visualizations
        visualizer = MapVisualizer(config)
        
        # Extract geo data from report
        content = report_data.get('full_content', '')
        geo_data = visualizer.extract_geo_data_from_report(content)
        
        # Province map (fallback sample data)
        province_data = {
            report.city or "四川": 100,
            "北京": 80,
            "上海": 90,
            "广东": 85
        }
        province_map = visualizer.generate_province_map(
            province_data, 
            title=f"{report.industry}产业分布地图"
        )
        
        # Geo scatter
        geo_scatter = visualizer.generate_geo_scatter(
            geo_data.get('cities', []),
            title=f"{report.city or ''}产业地理分布"
        )
        
        # 3D bar (mock using keywords)
        bar3d_data = [
            {"x": "市场规模", "y": report.city or "成都", "z": 100},
            {"x": "投资", "y": report.city or "成都", "z": 80},
            {"x": "企业数", "y": report.city or "成都", "z": 60}
        ]
        bar_3d = visualizer.generate_3d_bar_chart(bar3d_data)
        
        # Simple network graph
        nodes = [
            {"id": "1", "name": report.city or "城市", "category": 0, "value": 100},
            {"id": "2", "name": report.industry or "产业", "category": 1, "value": 80},
            {"id": "3", "name": "龙头企业", "category": 2, "value": 60}
        ]
        links = [
            {"source": "1", "target": "2", "value": 1},
            {"source": "2", "target": "3", "value": 1}
        ]
        network_graph = visualizer.generate_industry_network(nodes, links)
        
        return jsonify({
            'province_map': province_map,
            'bar_3d': bar_3d,
            'geo_scatter': geo_scatter,
            'network_graph': network_graph
        })
    
    except Exception as e:
        logger.error(f"Error generating visualizations: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/trend-analysis', methods=['POST'])
@login_required
def api_trend_analysis():
    """API endpoint for trend analysis of multiple reports."""
    try:
        data = request.get_json()
        report_ids = data.get('report_ids', [])
        metric = data.get('metric', 'market_size')
        
        if len(report_ids) < 3:
            return jsonify({'error': '至少需要3份报告进行趋势分析'}), 400
        
        analyzer = TrendAnalyzer()
        
        for report_id in report_ids:
            report = Report.query.filter_by(report_id=report_id).first()
            if not report or (report.user_id != current_user.id and current_user.role != 'admin'):
                continue
            file_path = _resolve_report_file_path(report_id, report)
            if not file_path:
                continue
            with open(file_path, 'r', encoding='utf-8') as f:
                report_data = json.load(f)
            
            time_str = analyzer.extract_time_from_report(
                report_data.get('full_content', ''),
                filename=report.file_path
            ) or report.created_at.strftime("%Y-%m-%d")
            
            analyzer.add_historical_report(report_id, report_data, time_str)
        
        chart_config = analyzer.generate_trend_chart_config(
            metric,
            include_prediction=True,
            prediction_periods=6
        )
        trend = analyzer.calculate_trend(metric)
        prediction = analyzer.predict_future(metric, 6)
        
        # Map to frontend fields
        resp = {
            'chart': chart_config,
            'trend_direction': trend.get('trend_direction'),
            'avg_growth_rate': trend.get('avg_growth_rate'),
            'data_points': trend.get('data_points'),
            'prediction_periods': prediction.get('prediction_periods', 6)
        }
        return jsonify(resp)
    
    except Exception as e:
        logger.error(f"Error in trend analysis: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/comparison', methods=['POST'])
@login_required
def api_comparison_analysis():
    """API endpoint for multi-document comparison."""
    try:
        data = request.get_json()
        report_ids = data.get('report_ids', [])
        
        if len(report_ids) < 2:
            return jsonify({'error': '至少需要2份报告进行对比分析'}), 400
        
        # Initialize analyzer
        analyzer = ComparisonAnalyzer()
        
        valid_count = 0
        # Load and add reports
        for report_id in report_ids:
            report = Report.query.filter_by(report_id=report_id).first()
            if not report or (report.user_id != current_user.id and current_user.role != 'admin'):
                continue
            
            # Load report data with robust path handling
            try:
                file_path = _resolve_report_file_path(report_id, report)
                if not file_path:
                    raise FileNotFoundError('报告文件不存在')
                with open(file_path, 'r', encoding='utf-8') as f:
                    report_data = json.load(f)
                
                # Add to comparison
                metadata = {
                    "name": f"{report.city}-{report.industry}" if report.city and report.industry else report.title or report.report_id,
                    "city": report.city,
                    "industry": report.industry
                }
                analyzer.add_report(report_id, report_data, metadata)
                valid_count += 1
            except Exception as e:
                logger.error(f"Error loading report {report_id} for comparison: {e}")
                continue
        
        if valid_count < 2:
            return jsonify({'error': '有效报告不足2份，无法进行对比'}), 400
        
        # Generate comparison analysis
        comparison = analyzer.compare_reports()
        text_report = analyzer.generate_comparison_report()
        
        # Prepare ranking (frontend expects `ranking`)
        ranking = []
        try:
            overall = comparison.get('rankings', {}).get('overall', [])
            ranking = [{
                'name': item.get('report_name', item.get('report_id')),
                'score': item.get('score', 0)
            } for item in overall]
        except Exception:
            ranking = []
        
        # Choose a default bar chart metric to present (market_size preferred)
        bar_chart = {}
        metrics = list(comparison.get('metric_comparison', {}).keys())
        preferred = 'market_size' if 'market_size' in metrics else (metrics[0] if metrics else None)
        if preferred:
            bar_chart = analyzer.generate_comparison_chart(preferred)
        
        return jsonify({
            'ranking': ranking,
            'bar_chart': bar_chart,
            'text_report': text_report
        })
    
    except Exception as e:
        logger.error(f"Error in comparison analysis: {e}")
        return jsonify({'error': str(e)}), 500


def setup_directories():
    """Ensure all necessary directories exist."""
    directories = [
        'data/input',
        'data/output',
        'data/output/llm_reports',
        'static/css',
        'static/js',
        'templates'
    ]
    
    for directory in directories:
        os.makedirs(directory, exist_ok=True)


@app.route('/stream-report/<task_id>')
@login_required
def stream_report_page(task_id):
    """Stream report generation page with dedicated URL"""
    try:
        # Get parameters from query string or form
        city = request.args.get('city', '')
        industry = request.args.get('industry', '')
        llm_service = request.args.get('llm_service', 'kimi')
        
        # If parameters not in query string, try to get from database
        if not city or not industry:
            # Try to find report by task_id (this is a best-effort fallback)
            # In production, you might want to store task_id in Report model
            logger.warning(f"Missing city/industry for task {task_id}, using defaults")
            city = city or '未知城市'
            industry = industry or '未知行业'
        
        return render_template('stream_report.html',
                             task_id=task_id,
                             city=city,
                             industry=industry,
                             llm_service=llm_service)
    except Exception as e:
        logger.error(f"Error loading stream report page: {e}")
        flash('无法加载报告生成页面', 'error')
        return redirect(url_for('index'))

@app.route('/settings')
@login_required
def settings():
    """Configuration settings page."""
    try:
        config_path = 'config.json'
        config = {}
        
        if os.path.exists(config_path):
            with open(config_path, 'r', encoding='utf-8') as f:
                config = json.load(f)
        
        return render_template('settings.html', config=config)
        
    except Exception as e:
        logger.error(f"Error loading settings: {e}")
        return render_template('settings.html', config={})

@app.route('/ui-states')
@login_required
def ui_states_page():
    return render_template('ui_states.html')


@app.route('/reports')
@login_required
def reports_list():
    """Display all reports for current user."""
    try:
        # Only show usable (completed + file exists) reports, separated by type
        base_query = Report.query
        if current_user.role != 'admin':
            base_query = base_query.filter_by(user_id=current_user.id)
        all_reports = base_query.order_by(Report.created_at.desc()).all()

        usable_llm = []
        usable_upload = []
        for r in all_reports:
            if r.status != 'completed':
                continue
            # Verify file exists
            try:
                if r.report_type == 'llm':
                    default_path = APP_ROOT / 'data' / 'output' / 'llm_reports' / f"{r.report_id}.json"
                    file_path = default_path if default_path.exists() else Path(r.file_path)
                else:
                    file_path = Path(r.file_path)
                if not file_path.exists():
                    continue
            except Exception:
                continue
            if r.report_type == 'llm':
                usable_llm.append(r)
            else:
                usable_upload.append(r)
        
        return render_template('reports_list.html', reports_llm=usable_llm, reports_upload=usable_upload)
    except Exception as e:
        logger.error(f"Error loading reports: {e}")
        flash('无法加载报告列表', 'error')
        return redirect(url_for('index'))


@app.route('/comparison')
@login_required  
def comparison_page():
    """Multi-document comparison page."""
    try:
        base_q = Report.query
        if current_user.role != 'admin':
            base_q = base_q.filter_by(user_id=current_user.id)
        candidates = base_q.filter_by(status='completed').order_by(Report.created_at.desc()).all()
        valid_reports = []
        for r in candidates:
            if _resolve_report_file_path(r.report_id, r):
                valid_reports.append(r)
        return render_template('comparison.html', reports=valid_reports)
    except Exception as e:
        logger.error(f"Error loading comparison page: {e}")
        flash('无法加载对比分析页面', 'error')
        return redirect(url_for('index'))


@app.route('/trend-analysis')
@login_required
def trend_analysis_page():
    """Trend analysis page."""
    try:
        base_q = Report.query
        if current_user.role != 'admin':
            base_q = base_q.filter_by(user_id=current_user.id)
        candidates = base_q.filter_by(status='completed').order_by(Report.created_at.desc()).all()
        valid_reports = [r for r in candidates if _resolve_report_file_path(r.report_id, r)]
        return render_template('trend_analysis.html', reports=valid_reports)
    except Exception as e:
        logger.error(f"Error loading trend analysis page: {e}")
        flash('无法加载趋势分析页面', 'error')
        return redirect(url_for('index'))


@app.route('/v3_roadmap')
def v3_roadmap():
    """Display the V3.0 roadmap page."""
    return render_template('v3_roadmap.html')


@app.route('/knowledge-graph')
@login_required
def knowledge_graph_page():
    """Knowledge graph page."""
    try:
        base_q = Report.query.filter_by(report_type='llm', status='completed')
        if current_user.role != 'admin':
            base_q = base_q.filter_by(user_id=current_user.id)
        candidates = base_q.order_by(Report.created_at.desc()).limit(20).all()
        valid_reports = [r for r in candidates if _resolve_report_file_path(r.report_id, r)]
        return render_template('knowledge_graph.html', reports=valid_reports)
    except Exception as e:
        logger.error(f"Error loading knowledge graph page: {e}")
        flash('无法加载知识图谱页面', 'error')
        return redirect(url_for('index'))


@app.route('/map-visualization')
@login_required
def map_visualization_page():
    """Map visualization page."""
    try:
        base_q = Report.query.filter_by(status='completed')
        if current_user.role != 'admin':
            base_q = base_q.filter_by(user_id=current_user.id)
        candidates = base_q.order_by(Report.created_at.desc()).limit(20).all()
        valid_reports = [r for r in candidates if _resolve_report_file_path(r.report_id, r)]
        return render_template('map_visualization.html', reports=valid_reports)
    except Exception as e:
        logger.error(f"Error loading map visualization page: {e}")
        flash('无法加载地图可视化页面', 'error')
        return redirect(url_for('index'))


@app.route('/poi-map-visualization')
@login_required
def poi_map_visualization_page():
    """POI map visualization page for uploading and visualizing enterprise/university/research institutes data."""
    try:
        # Load config for Baidu Map AK
        config_path = 'config.json'
        config = {}
        baidu_ak = ''
        google_map_key = ''
        if os.path.exists(config_path):
            with open(config_path, 'r', encoding='utf-8') as f:
                config = json.load(f)
                baidu_ak = config.get('api_keys', {}).get('baidu_map', '')
                google_map_key = config.get('api_keys', {}).get('google_map', '')

        return render_template('poi_map_visualization.html', 
                             config=config,
                             baidu_ak=baidu_ak,
                             google_map_key=google_map_key)
    except Exception as e:
        logger.error(f"Error loading POI map visualization page: {e}")
        flash('无法加载POI地图可视化页面', 'error')
        return redirect(url_for('index'))


@app.route('/industry-chain')
@login_required
def industry_chain_page():
    """Industry chain analysis page."""
    try:
        base_q = Report.query.filter_by(status='completed')
        if current_user.role != 'admin':
            base_q = base_q.filter_by(user_id=current_user.id)
        candidates = base_q.order_by(Report.created_at.desc()).limit(20).all()
        valid_reports = [r for r in candidates if _resolve_report_file_path(r.report_id, r)]
        return render_template('industry_chain.html', reports=valid_reports)
    except Exception as e:
        logger.error(f"Error loading industry chain page: {e}")
        flash('无法加载产业链分析页面', 'error')
        return redirect(url_for('index'))


@app.route('/casual-learning')
@login_required
def casual_learning_page():
    """Casual learning page for regional economic information."""
    try:
        # Render the casual learning page which contains detailed information about the selected regions
        return render_template('casual_learning.html')
    except Exception as e:
        logger.error(f"Error loading casual learning page: {e}")
        flash('无法加载顺手了解页面', 'error')
        return redirect(url_for('index'))


@app.route('/api/leadership/<region>')
@login_required
def get_leadership_data(region):
    """API endpoint to get leadership data for a specific region."""
    try:
        # Get leadership data using the scraper
        leadership_data = leadership_scraper.get_leadership_data(region)

        return jsonify({
            'success': True,
            'data': leadership_data,
            'timestamp': now_beijing().strftime('%Y年%m月%d日')
        })
    except Exception as e:
        logger.error(f"Error getting leadership data for region {region}: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/policy-analysis')
@login_required
def policy_analysis_page():
    """Policy analysis page."""
    return render_template('policy_analysis.html')

# Policy search and aggregation (dual-dimension)
@app.route('/api/policy/search', methods=['POST'])
@login_required
def api_policy_search():
    try:
        data = request.get_json() or {}
        region = (data.get('region') or '').strip()  # e.g., 四川/成都/高新区
        region_code = (data.get('region_code') or '${placeholder_region_code}').strip()
        industry_tags = data.get('industry_tags') or []  # e.g., ["人工智能", "AI芯片"]
        years = int(data.get('time_range_years') or 3)

        if not region or not industry_tags:
            return jsonify({'error': '请输入区域与产业标签'}), 400

        # Caching: Redis (hot), file+memory (fallback)
        cache_key = f"policy_search:{region}:{region_code}:{','.join(industry_tags)}:{years}"
        cached = None
        # Try Redis first
        try:
            import os
            import redis
            redis_url = os.getenv('REDIS_URL')
            if redis_url:
                r = redis.Redis.from_url(redis_url)
                hit = r.get(cache_key)
                if hit:
                    cached = json.loads(hit)
        except Exception:
            cached = None
        # Fallback local cache
        if not cached:
            cm = CacheManager(ttl=3600)
            cached = cm.get(cache_key)
        if cached:
            return jsonify({'success': True, 'cached': True, 'data': cached})

        # Aggregate policies from government sources
        from src.data.web_scraper import WebScraper
        scraper = WebScraper(timeout=20)
        agg = scraper.aggregate_policies(region_name=region, region_code=region_code, industry_tags=industry_tags, years=years)
        scraper.close()

        # Fetch WeChat articles and merge with government policies
        try:
            # Get WeChat articles from database that match the region and industry
            from sqlalchemy import or_, and_

            # Search for WeChat articles matching the region or industry
            wechat_articles = []
            for industry_tag in industry_tags:
                # Search by industry relevance and region
                articles_query = WeChatArticle.query.filter(
                    or_(
                        WeChatArticle.industry_relevance.like(f'%{industry_tag}%'),
                        WeChatArticle.title.like(f'%{industry_tag}%'),
                        WeChatArticle.content.like(f'%{industry_tag}%')
                    )
                )

                # Filter by region if specified
                if region:
                    articles_query = articles_query.filter(
                        or_(
                            WeChatArticle.source_account.like(f'%{region}%'),
                            # Note: WeChatArticle doesn't currently have a region field,
                            # but when the task runs it includes region in industry_relevance or source_account
                        )
                    )

                # Filter by time range
                from datetime import datetime, timedelta
                time_limit = now_beijing() - timedelta(days=years*365)
                articles_query = articles_query.filter(
                    WeChatArticle.created_at >= time_limit
                )

                articles = articles_query.all()

                for article in articles:
                    wechat_articles.append({
                        'title': article.title,
                        'publish_date': article.publish_date,
                        'source': article.source_account,
                        'summary': article.summary,
                        'url': article.url,
                        'type': '微信公众号文章',
                        'region_code': region_code,  # Use provided region code
                        'industry_tags': industry_tags,
                        'official_link': article.url,
                        'timeliness': '有效' if article.publish_date else '未知'
                    })

            # Add WeChat articles to the government policies
            agg.get('policies', []).extend(wechat_articles)

            # Update summary with WeChat articles count
            current_wechat_count = agg['summary'].get('wechat_count', 0)
            agg['summary']['wechat_count'] = current_wechat_count + len(wechat_articles)
            agg['summary']['total_count'] = agg['summary'].get('count', 0) + len(wechat_articles)

            # Update timeline with WeChat articles
            for article in wechat_articles:
                agg['timeline'].append({
                    'date': article.get('publish_date', ''),
                    'type': '微信文章',
                    'title': article.get('title', ''),
                    'source': article.get('source', '')
                })

        except Exception as e:
            logger.error(f"Error fetching WeChat articles: {e}")
            # Continue without WeChat articles if there's an error

        # Save to caches
        try:
            if redis_url:
                r = redis.Redis.from_url(redis_url)
                r.setex(cache_key, 3600, json.dumps(agg, ensure_ascii=False))
        except Exception:
            pass
        cm = CacheManager(ttl=3600)
        cm.set(cache_key, agg)

        return jsonify({'success': True, 'cached': False, 'data': agg})
    except Exception as e:
        logger.error(f"Error in policy search: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return jsonify({'error': f'政策检索失败: {str(e)}'}), 500

@app.route('/api/policy/export', methods=['POST'])
@login_required
def api_policy_export():
    try:
        data = request.get_json() or {}
        region = (data.get('region') or '').strip()
        region_code = (data.get('region_code') or '${placeholder_region_code}').strip()
        industry_tags = data.get('industry_tags') or []
        years = int(data.get('time_range_years') or 3)

        if not region or not industry_tags:
            return jsonify({'error': '请输入区域与产业标签'}), 400

        # Reuse search aggregation (prefer cache)
        req = {
            'region': region,
            'region_code': region_code,
            'industry_tags': industry_tags,
            'time_range_years': years
        }
        with app.test_request_context('/api/policy/search', method='POST', json=req):
            search_resp = api_policy_search().json
        agg = (search_resp or {}).get('data') or {}

        # Compose export-friendly report_data
        title = f"{region} {industry_tags[0] if industry_tags else ''} 产业政策包"
        report_data = {
            'title': title,
            'city': region,
            'industry': industry_tags[0] if industry_tags else '',
            'generated_at': format_beijing_time(None, '%Y-%m-%d %H:%M'),
            'summary': {
                'zh': f"共收录近{years}年政策 {agg.get('summary', {}).get('count', 0)} 条",
            },
            'sections': {
                'policy_landscape': '\n'.join([
                    f"- {p.get('publish_date')}｜{p.get('source')}｜{p.get('title')}"
                    for p in agg.get('policies', [])
                ])
            }
        }

        exporter = ReportExporter()
        safe_name = f"policy_package_{region}_{industry_tags[0] if industry_tags else 'industry'}_{now_beijing().strftime('%Y%m%d')}"
        pdf_path = exporter.export_to_pdf(report_data, safe_name)
        excel_path = exporter.export_to_excel(report_data, safe_name)

        return jsonify({'success': True, 'export_paths': {'pdf': pdf_path, 'excel': excel_path}})
    except Exception as e:
        logger.error(f"Error in policy export: {e}")
        return jsonify({'error': f'导出失败: {str(e)}'}), 500

@app.route('/api/policy/content', methods=['POST'])
@login_required
def api_policy_content():
    try:
        data = request.get_json() or {}
        url = (data.get('url') or '').strip()
        title = (data.get('title') or '').strip()
        source = (data.get('source') or '').strip()
        publish_date = (data.get('publish_date') or '').strip()
        if not url and not title:
            return jsonify({'error': '缺少内容标识'}), 400
        sample_html = (
            f"<h2>{title or '产业支持政策'}</h2>"+
            f"<p>发布日期：{publish_date or '2024-12-01'}｜发布机构：{source or '市政府办公厅'}</p>"+
            "<h3>一、总体要求</h3>"
            "<p>为促进产业高质量发展，围绕创新、资金、人才、载体等重点方向提出系统性举措。</p>"
            "<h3>二、支持措施</h3>"
            "<ol><li>资金补助：对重点项目给予最高500万元补贴。</li><li>税收优惠：对高新技术企业减按15%征收企业所得税。</li><li>人才引进：提供安家补贴与住房保障。</li></ol>"
            "<h3>三、申报流程</h3>"
            "<table class='table table-bordered'><thead><tr><th>环节</th><th>时间</th><th>说明</th></tr></thead><tbody><tr><td>材料提交</td><td>每季度</td><td>通过线上系统提交</td></tr><tr><td>资格审核</td><td>15个工作日</td><td>部门联合评审</td></tr><tr><td>资金拨付</td><td>公告后</td><td>分期拨付</td></tr></tbody></table>"
            "<h3>四、监督管理</h3>"
            "<p>建立绩效评估与动态调整机制，强化资金使用监管。</p>"
        )
        import re
        heads = re.findall(r'<h(\d)>(.*?)</h\1>', sample_html)
        headings = [{'level': int(lvl), 'text': txt} for lvl, txt in heads]
        plain = re.sub(r'<[^>]+>', '', sample_html)
        chunks = []
        step = 1200
        for i in range(0, len(plain), step):
            chunks.append(sample_html)
        download_url = ''
        return jsonify({
            'success': True,
            'meta': {
                'title': title or '产业支持政策',
                'source': source or '市政府办公厅',
                'publish_date': publish_date or '2024-12-01',
                'url': url
            },
            'headings': headings,
            'content_html': sample_html,
            'content_pages': chunks,
            'download_url': download_url
        })
    except Exception as e:
        logger.error(f"Error in policy content: {e}")
        return jsonify({'error': f'获取正文失败: {str(e)}'}), 500

@app.route('/api/data-sources', methods=['GET'])
@login_required
def api_data_sources():
    try:
        sources = [
            {
                'provider': '国务院',
                'name': '国务院政策文件库',
                'url': 'https://www.gov.cn/zhengce/',
                'format': ['HTML'],
                'last_update': format_beijing_time(None, '%Y-%m-%d %H:%M'),
                'verified': True,
                'quality': {'completeness': 0.95, 'accuracy': 0.97, 'freshness': 0.92},
                'https': True
            },
            {
                'provider': '国家发展改革委',
                'name': '发改委政策发布',
                'url': 'https://www.ndrc.gov.cn/xxgk/ztzl/',
                'format': ['HTML'],
                'last_update': format_beijing_time(None, '%Y-%m-%d %H:%M'),
                'verified': True,
                'quality': {'completeness': 0.9, 'accuracy': 0.95, 'freshness': 0.9},
                'https': True
            },
            {
                'provider': '工业和信息化部',
                'name': '工信部政策发布',
                'url': 'https://www.miit.gov.cn/gzcy/zcwj/',
                'format': ['HTML'],
                'last_update': format_beijing_time(None, '%Y-%m-%d %H:%M'),
                'verified': True,
                'quality': {'completeness': 0.88, 'accuracy': 0.94, 'freshness': 0.9},
                'https': True
            },
            {
                'provider': '四川省人民政府',
                'name': '四川省政府政策文件库',
                'url': 'https://www.sc.gov.cn/zwgk/tzgg/',
                'format': ['HTML'],
                'last_update': format_beijing_time(None, '%Y-%m-%d %H:%M'),
                'verified': True,
                'quality': {'completeness': 0.92, 'accuracy': 0.94, 'freshness': 0.95},
                'https': True
            },
            {
                'provider': '四川省科学技术厅',
                'name': '四川省科技厅政策发布',
                'url': 'http://kjt.sc.gov.cn/kjt/c100382/xxgk.shtml',
                'format': ['HTML'],
                'last_update': format_beijing_time(None, '%Y-%m-%d %H:%M'),
                'verified': True,
                'quality': {'completeness': 0.89, 'accuracy': 0.93, 'freshness': 0.96},
                'https': True
            }
        ]
        return jsonify({'success': True, 'sources': sources})
    except Exception as e:
        logger.error(f"Error in data sources: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/wechat-articles', methods=['POST'])
@login_required
def api_wechat_articles():
    """API endpoint to get WeChat articles by region and industry"""
    try:
        data = request.get_json() or {}
        region = (data.get('region') or '').strip()
        industry = (data.get('industry') or '').strip()
        limit = int(data.get('limit', 10))

        if not region or not industry:
            return jsonify({'error': '请输入区域与产业名称'}), 400

        # Get WeChat articles
        scraper = WeChatScraper()
        articles = scraper.get_articles_by_region_and_industry(region, industry, limit)
        scraper.close()

        # Convert to serializable format
        serialized_articles = []
        for article in articles:
            serialized_articles.append({
                'title': article.title,
                'content': article.content,
                'publish_date': article.publish_date,
                'author': article.author,
                'source_account': article.source_account,
                'url': article.url,
                'summary': article.summary,
                'keywords': article.keywords,
                'industry_relevance': article.industry_relevance
            })

        return jsonify({
            'success': True,
            'articles': serialized_articles,
            'count': len(serialized_articles)
        })
    except Exception as e:
        logger.error(f"Error in WeChat articles API: {e}")
        return jsonify({'error': f'获取微信公众号文章失败: {str(e)}'}), 500

@app.route('/api/wechat-search', methods=['POST'])
@login_required
def api_wechat_search():
    """API endpoint to search WeChat articles by keywords"""
    try:
        data = request.get_json() or {}
        keywords = data.get('keywords', [])
        region = (data.get('region') or '').strip()
        limit = int(data.get('limit', 10))

        if not keywords:
            return jsonify({'error': '请输入搜索关键词'}), 400

        # Search WeChat articles
        scraper = WeChatScraper()
        articles = scraper.search_articles(keywords, region, limit)
        scraper.close()

        # Convert to serializable format
        serialized_articles = []
        for article in articles:
            serialized_articles.append({
                'title': article.title,
                'content': article.content,
                'publish_date': article.publish_date,
                'author': article.author,
                'source_account': article.source_account,
                'url': article.url,
                'summary': article.summary,
                'keywords': article.keywords,
                'industry_relevance': article.industry_relevance
            })

        return jsonify({
            'success': True,
            'articles': serialized_articles,
            'count': len(serialized_articles)
        })
    except Exception as e:
        logger.error(f"Error in WeChat search API: {e}")
        return jsonify({'error': f'搜索微信公众号文章失败: {str(e)}'}), 500

@app.route('/api/wechat-analyze', methods=['POST'])
@login_required
def api_wechat_analyze():
    """API endpoint to analyze WeChat article content"""
    try:
        data = request.get_json() or {}
        content = data.get('content', '')
        title = data.get('title', '')
        source_account = data.get('source_account', '')
        publish_date = data.get('publish_date', '')
        author = data.get('author', '')
        url = data.get('url', '')

        if not content:
            return jsonify({'error': '请输入文章内容'}), 400

        from src.data.wechat_scraper import WeChatArticle
        article = WeChatArticle(
            title=title,
            content=content,
            publish_date=publish_date,
            author=author,
            source_account=source_account,
            url=url
        )

        # Analyze the article
        analyzer = WeChatContentAnalyzer()
        result = analyzer.analyze_article_content(article)

        return jsonify({
            'success': True,
            'analysis': result
        })
    except Exception as e:
        logger.error(f"Error in WeChat analysis API: {e}")
        return jsonify({'error': f'分析微信公众号文章失败: {str(e)}'}), 500

@app.route('/api/wechat-account-articles', methods=['POST'])
@login_required
def api_wechat_account_articles():
    """API endpoint to get articles from a specific WeChat account"""
    try:
        data = request.get_json() or {}
        account_name = (data.get('account_name') or '').strip()
        limit = int(data.get('limit', 10))

        if not account_name:
            return jsonify({'error': '请输入公众号名称'}), 400

        # Get articles from specific account
        scraper = WeChatScraper()
        articles = scraper.get_account_articles(account_name, limit)
        scraper.close()

        # Convert to serializable format
        serialized_articles = []
        for article in articles:
            serialized_articles.append({
                'title': article.title,
                'content': article.content,
                'publish_date': article.publish_date,
                'author': article.author,
                'source_account': article.source_account,
                'url': article.url,
                'summary': article.summary,
                'keywords': article.keywords,
                'industry_relevance': article.industry_relevance
            })

        return jsonify({
            'success': True,
            'articles': serialized_articles,
            'count': len(serialized_articles)
        })
    except Exception as e:
        logger.error(f"Error in WeChat account articles API: {e}")
        return jsonify({'error': f'获取公众号文章失败: {str(e)}'}), 500

@app.route('/wechat-config')
@login_required
def wechat_config_page():
    """WeChat account configuration page."""
    return render_template('wechat_config.html')

@app.route('/api/links', methods=['GET'])
@login_required
def api_links_list():
    try:
        links = [
            {'url': 'https://www.gov.cn/zhengce/', 'type': '政策原文', 'importance': 'high'},
            {'url': 'https://www.sc.gov.cn/zwgk/tzgg/', 'type': '政策原文', 'importance': 'high'},
            {'url': 'https://apply.sckjt.net.cn/egrant/apply', 'type': '申报入口', 'importance': 'high'},
            {'url': 'https://kjt.sc.gov.cn/kjt/c100382/xxgk.shtml', 'type': '部门页面', 'importance': 'high'},
            {'url': 'https://www.ndrc.gov.cn/xxgk/zcjd/', 'type': '解释解读', 'importance': 'medium'},
            {'url': 'https://www.miit.gov.cn/gzcy/zcwj/', 'type': '政策原文', 'importance': 'medium'},
            {'url': 'https://www.cdst.gov.cn/cdst/c101523/xxgk.shtml', 'type': '部门页面', 'importance': 'medium'},
            {'url': 'https://zcykt.sc.gov.cn/', 'type': '政策原文', 'importance': 'low'}
        ]
        return jsonify({'success': True, 'links': links})
    except Exception as e:
        logger.error(f"Error in links list: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/link-check', methods=['POST'])
@login_required
def api_link_check():
    try:
        data = request.get_json() or {}
        items = data.get('links') or []
        results = []
        import requests
        for it in items:
            url = (it.get('url') or '').strip()
            if not url:
                results.append({'url': url, 'reachable': False, 'status': 'invalid', 'https': False, 'checked_at': format_beijing_time(None, '%Y-%m-%dT%H:%M:%S')})
                continue
            https = url.startswith('https://')
            status = 'error'
            reachable = False
            code = None
            try:
                r = requests.head(url, timeout=5, allow_redirects=True)
                code = r.status_code
                reachable = 200 <= r.status_code < 400
                status = 'ok' if reachable else f'http_{r.status_code}'
            except Exception:
                status = 'error'
            results.append({'url': url, 'reachable': reachable, 'status': status, 'https': https, 'code': code, 'type': it.get('type'), 'importance': it.get('importance'), 'checked_at': format_beijing_time(None, '%Y-%m-%dT%H:%M:%S')})
        return jsonify({'success': True, 'results': results})
    except Exception as e:
        logger.error(f"Error in link check: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/links/click', methods=['POST'])
@login_required
def api_link_click():
    try:
        data = request.get_json() or {}
        url = (data.get('url') or '').strip()
        if not url:
            return jsonify({'error': '缺少链接'}), 400
        cm = CacheManager(ttl=24 * 3600)
        key = f"link_stats:{url}"
        cur = cm.get(key) or {'count': 0}
        cur['count'] = int(cur.get('count', 0)) + 1
        cm.set(key, cur)
        return jsonify({'success': True, 'url': url, 'count': cur['count']})
    except Exception as e:
        logger.error(f"Error in link click: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/links/stats', methods=['GET'])
@login_required
def api_links_stats():
    try:
        urls = request.args.get('urls') or ''
        urls_list = [u.strip() for u in urls.split(',') if u.strip()]
        cm = CacheManager(ttl=24 * 3600)
        data = {}
        for u in urls_list:
            key = f"link_stats:{u}"
            cur = cm.get(key) or {'count': 0}
            data[u] = cur
        return jsonify({'success': True, 'stats': data})
    except Exception as e:
        logger.error(f"Error in links stats: {e}")
        return jsonify({'error': str(e)}), 500

# Terminology page and API
@app.route('/terminology')
@login_required
def terminology_page():
    return render_template('terminology.html')

@app.route('/api/terminology', methods=['GET', 'POST'])
@login_required
def api_terminology():
    try:
        path = APP_ROOT / 'data' / 'terminology.json'
        if request.method == 'GET':
            if path.exists():
                with open(path, 'r', encoding='utf-8') as f:
                    return jsonify(json.load(f))
            return jsonify({})
        else:
            data = request.get_json() or {}
            term = data.get('term')
            info = data.get('info', {})
            existing = {}
            if path.exists():
                with open(path, 'r', encoding='utf-8') as f:
                    existing = json.load(f)
            if term:
                existing[term] = info
                path.parent.mkdir(parents=True, exist_ok=True)
                with open(path, 'w', encoding='utf-8') as f:
                    json.dump(existing, f, ensure_ascii=False, indent=2)
            return jsonify({'success': True})
    except Exception as e:
        logger.error(f"Error in terminology API: {e}")
        return jsonify({'error': str(e)}), 500

# Knowledge graph API
@app.route('/api/report/<report_id>/knowledge-graph')
@login_required
def api_knowledge_graph(report_id):
    try:
        report = Report.query.filter_by(report_id=report_id).first()
        if not report or (report.user_id != current_user.id and current_user.role != 'admin'):
            return jsonify({'error': '无权限访问'}), 403
        # Load report
        if report.report_type == 'llm':
            default_path = APP_ROOT / 'data' / 'output' / 'llm_reports' / f"{report_id}.json"
            file_path = default_path if default_path.exists() else Path(report.file_path)
        else:
            file_path = Path(report.file_path)
        with open(file_path, 'r', encoding='utf-8') as f:
            report_data = json.load(f)
        # Extract entities
        extractor = EntityExtractor()
        entities = extractor.extract_entities(report_data.get('full_content', ''))
        # Build graph
        kg = KnowledgeGraphVisualizer()
        graph = kg.transform_entities_to_graph(entities)
        config = kg.generate_echarts_config(graph, title=f"{report.city or ''}{report.industry or ''} 知识图谱")
        return jsonify({
            'graph_config': config,
            'node_count': graph.get('statistics', {}).get('total_nodes', 0),
            'edge_count': graph.get('statistics', {}).get('total_links', 0),
            'type_count': graph.get('statistics', {}).get('node_types', 0)
        })
    except Exception as e:
        logger.error(f"Error in knowledge graph: {e}")
        return jsonify({'error': str(e)}), 500

# Industry chain API
@app.route('/api/report/<report_id>/industry-chain')
@login_required
def api_industry_chain(report_id):
    try:
        report = Report.query.filter_by(report_id=report_id).first()
        if not report or (report.user_id != current_user.id and current_user.role != 'admin'):
            return jsonify({'error': '无权限访问'}), 403
        # Load report
        if report.report_type == 'llm':
            default_path = APP_ROOT / 'data' / 'output' / 'llm_reports' / f"{report_id}.json"
            file_path = default_path if default_path.exists() else Path(report.file_path)
        else:
            file_path = Path(report.file_path)
        with open(file_path, 'r', encoding='utf-8') as f:
            report_data = json.load(f)
        text = report_data.get('sections', {}).get('value_chain') or report_data.get('full_content', '')
        # Simple heuristic parsing
        def section(items):
            return {
                'strength': '良好',
                'components': items[:6],
                'gaps': items[6:9]
            }
        # Extract bullet-like lines
        import re
        items = [ln.strip('- •* ').strip() for ln in text.split('\n') if re.match(r'^\s*[-•*]|\d+[\.\)、]', ln)]
        upstream = section(items[:9])
        midstream = section(items[3:12])
        downstream = section(items[6:15])
        weak_points = [it for it in items if any(k in it for k in ['瓶颈', '缺', '不足', '短板'])][:5]
        return jsonify({
            'completeness_score': 78,
            'upstream': upstream,
            'midstream': midstream,
            'downstream': downstream,
            'weak_points': weak_points
        })
    except Exception as e:
        logger.error(f"Error in industry chain: {e}")
        return jsonify({'error': str(e)}), 500

# Policy analysis API
@app.route('/api/policy-analysis', methods=['POST'])
@login_required
def api_policy_analysis():
    try:
        if 'file' not in request.files:
            return jsonify({'error': '未上传文件'}), 400
        file = request.files['file']
        filename = file.filename
        ext = filename.rsplit('.', 1)[-1].lower()
        content = ''

        # File parsing based on format - need to save to temp and reopen for each format
        import tempfile
        import os
        from werkzeug.utils import secure_filename

        with tempfile.NamedTemporaryFile(delete=False, suffix=f'.{ext}') as temp_file:
            file.save(temp_file.name)
            temp_path = temp_file.name

        try:
            # File parsing based on format
            if ext in ['txt', 'md']:
                with open(temp_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
            elif ext in ['doc', 'docx']:
                if ext == 'docx':
                    # Handle .docx files
                    from docx import Document
                    doc = Document(temp_path)
                    content = '\n'.join(p.text for p in doc.paragraphs)
                else:  # ext == 'doc'
                    # For older .doc files, use mammoth as primary approach
                    try:
                        import mammoth
                        with open(temp_path, "rb") as doc_file:
                            result = mammoth.convert_to_text(doc_file)
                            content = result.value
                    except ImportError:
                        # Mammoth not available, try docx2txt as fallback (though mainly for docx)
                        try:
                            import docx2txt
                            content = docx2txt.process(temp_path)
                        except:
                            # Fallback to using docx library (which might work for some doc files)
                            try:
                                import subprocess
                                import os
                                # Try system utility antiword if available
                                if os.name != 'nt':  # Not Windows
                                    result = subprocess.run(['antiword', temp_path],
                                                          capture_output=True, text=True, timeout=30)
                                    if result.returncode == 0:
                                        content = result.stdout
                                    else:
                                        # If antiword not available, final fallback
                                        content = "DOC文件解析失败，请尝试转换为DOCX格式后重新上传。"
                                else:
                                    # On Windows, inform about limitation
                                    content = "DOC文件解析失败，请尝试转换为DOCX格式后重新上传。"
                            except FileNotFoundError:
                                # antiword not found, fall back to error message
                                content = "DOC文件解析失败，请尝试转换为DOCX格式后重新上传。"
                            except:
                                # Any other error in fallback chain
                                content = "DOC文件解析失败，请尝试转换为DOCX格式后重新上传。"
            elif ext == 'pdf':
                try:
                    import PyPDF2
                    with open(temp_path, 'rb') as f:
                        reader = PyPDF2.PdfReader(f)
                        for page in reader.pages:
                            content += page.extract_text() or ''
                except Exception:
                    # Fallback to pypdf if PyPDF2 fails
                    try:
                        import pypdf
                        with open(temp_path, 'rb') as f:
                            reader = pypdf.PdfReader(f)
                            for page in reader.pages:
                                content += page.extract_text() or ''
                    except Exception:
                        content = "PDF文件解析失败，请检查文件格式是否正确。"
            else:
                return jsonify({'error': '不支持的文件格式'}), 400
        finally:
            # Clean up temp file
            os.unlink(temp_path)

        # Get LLM service selection from request
        data = request.form if request.form else request.get_json()
        llm_service = 'kimi'  # Default to Kimi as requested
        if data and 'llm_service' in data:
            llm_service = data['llm_service']

        # Use hybrid parser with selected LLM service
        from hybrid_parser import HybridPolicyParser
        parser = HybridPolicyParser()
        structured_content = parser.parse_policy_document(content, llm_service=llm_service)

        # Generate visualization data based on structured content
        timeline_items = structured_content.get('timeline', [])
        visualization_data = {
            'timeline_chart': {
                'dates': [item.get('date', '') for item in timeline_items],
                'events': [item.get('event', item.get('description', '')) for item in timeline_items]
            },
            'industry_network': {
                'nodes': [{'name': ind, 'category': 0} for ind in structured_content['metadata'].get('key_industries', [])[:5]],
                'links': [{'source': i, 'target': i+1} for i in range(len(structured_content['metadata'].get('key_industries', [])[:4]))]
            },
            'heatmap_data': {
                'regions': structured_content['metadata'].get('applicable_regions', []),
                'intensity': [80, 60, 75, 90] if structured_content['metadata'].get('applicable_regions') else []
            },
            'radar_chart': {
                'dimensions': ['资金支持', '税收优惠', '人才支持', '技术扶持', '市场准入'],
                'values': [
                    len(structured_content['quantitative_data'].get('amounts', []))*20,
                    len([item for item in structured_content['quantitative_data'].get('ratios', []) if '%' in str(item) or '百分' in str(item)])*25,
                    70 if '人才' in content else 30,
                    80 if '技术' in content else 40,
                    60 if '市场' in content else 50
                ]
            }
        }

        return jsonify({
            'policy_metadata': structured_content['metadata'],
            'document_structure': structured_content.get('document_structure', []),
            'entities': structured_content.get('entities', {}),
            'provisions': structured_content['provisions'],
            'requirements': structured_content['requirements'],
            'quantitative_data': structured_content['quantitative_data'],
            'timeline': structured_content['timeline'],
            'relationships': structured_content.get('relationships', []),
            'key_points': structured_content['key_points'],
            'analysis': structured_content['analysis'],
            'visualization_data': visualization_data,
            # Removed applicability_score as per requirements
            'full_text': structured_content['full_text'],
            'llm_service_used': llm_service,  # Include the service used for debugging
            'official_link': ''
        })
    except Exception as e:
        logger.error(f"Error in policy analysis: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return jsonify({'error': str(e)}), 500

# Data story endpoints
@app.route('/data-story')
@login_required
def data_story_page():
    try:
        if current_user.role == 'admin':
            reports = Report.query.filter_by(status='completed').order_by(Report.created_at.desc()).limit(12).all()
        else:
            reports = Report.query.filter_by(user_id=current_user.id, status='completed').order_by(Report.created_at.desc()).limit(12).all()
        return render_template('story.html', reports=reports)
    except Exception:
        return render_template('story.html', reports=[])

@app.route('/api/report/<report_id>/story')
@login_required
def api_report_story(report_id):
    try:
        report = Report.query.filter_by(report_id=report_id).first()
        if not report or (report.user_id != current_user.id and current_user.role != 'admin'):
            return jsonify({'error': '无权限访问'}), 403
        
        # Use robust path resolution
        file_path = _resolve_report_file_path(report_id, report)
        if not file_path:
            return jsonify({'error': '报告文件不存在'}), 404
        
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        story_gen = StoryGenerator()
        # Get story type from query parameter, default to 'industry_overview'
        story_type = request.args.get('type', 'industry_overview')
        story = story_gen.create_story(data, story_type=story_type)
        return jsonify(story)
    except Exception as e:
        logger.error(f"Error generating story: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/report/<report_id>/story-view')
@login_required
def report_story_view(report_id):
    """Render data storytelling page for a report."""
    try:
        report = Report.query.filter_by(report_id=report_id).first()
        if not report or (report.user_id != current_user.id and current_user.role != 'admin'):
            flash('无权限访问该报告')
            return redirect(url_for('index'))
        
        # Use robust path resolution
        file_path = _resolve_report_file_path(report_id, report)
        if not file_path:
            flash('报告文件不存在')
            return redirect(url_for('index'))
        
        with open(file_path, 'r', encoding='utf-8') as f:
            report_data = json.load(f)
        
        # Generate story with type from query parameter
        story_gen = StoryGenerator()
        story_type = request.args.get('type', 'industry_overview')
        story = story_gen.create_story(report_data, story_type=story_type)
        
        return render_template('story_view.html', report=report, story=story, report_data=report_data)
    except Exception as e:
        logger.error(f"Error rendering story view: {e}")
        flash(f'生成数据故事失败: {str(e)}')
        return redirect(url_for('index'))

# Export API
@app.route('/api/export-report/<report_id>/<format>', endpoint='api_export_report')
@login_required
def api_export_report(report_id, format):
    """Export report to PDF, Word, or Excel."""
    try:
        from src.export.report_exporter import ReportExporter
        from flask import send_file
        
        report = Report.query.filter_by(report_id=report_id).first()
        if not report or (report.user_id != current_user.id and current_user.role != 'admin'):
            return jsonify({'error': '无权限访问'}), 403
        
        # Load report data
        file_path = _resolve_report_file_path(report_id, report)
        if not file_path:
            return jsonify({'error': '报告文件不存在'}), 404
        
        with open(file_path, 'r', encoding='utf-8') as f:
            report_data = json.load(f)
        
        # Normalize data for exporter
        report_data['city'] = report.city or report_data.get('metadata', {}).get('city') or 'N/A'
        report_data['industry'] = report.industry or report_data.get('metadata', {}).get('industry') or 'N/A'
        report_data['generated_at'] = report.created_at.strftime('%Y-%m-%d %H:%M')
        
        # If sections/full_content absent (e.g., uploaded analysis), compose sections from dashboard data
        if not report_data.get('sections') and not report_data.get('full_content'):
            sections = {}
            # Executive summary from key highlights or summary
            key_highlights = []
            if isinstance(report_data.get('summary'), dict):
                key_highlights = report_data['summary'].get('key_highlights', [])
            elif isinstance(report_data.get('key_insights'), list):
                key_highlights = [i.get('title') for i in report_data['key_insights'] if i.get('title')]
            if key_highlights:
                sections['executive_summary'] = '\n'.join(f"- {x}" for x in key_highlights)
            
            # Industry overview from categories
            if isinstance(report_data.get('categories'), dict):
                overview_lines = []
                for cat, cat_data in report_data['categories'].items():
                    pts = cat_data.get('key_points') or []
                    if pts:
                        overview_lines.append(f"【{cat}】")
                        for p in pts[:5]:
                            overview_lines.append(f"- {p}")
                if overview_lines:
                    sections['industry_overview'] = '\n'.join(overview_lines)
            
            # AI integration
            if isinstance(report_data.get('ai_opportunities'), dict):
                ai_lines = []
                for name, ai in report_data['ai_opportunities'].items():
                    rec = ai.get('recommendation', '')
                    score = ai.get('potential_score', 0)
                    ai_lines.append(f"- {name}: 潜力{score}/100; 建议: {rec}")
                if ai_lines:
                    sections['ai_integration'] = '\n'.join(ai_lines)
            
            # Statistics/conclusion
            if isinstance(report_data.get('statistics'), dict):
                stats = report_data['statistics']
                st_lines = [
                    f"总字符数: {stats.get('total_characters', '-')}",
                    f"总词数: {stats.get('total_words', '-')}",
                    f"段落数: {stats.get('total_segments', '-')}",
                    f"预计阅读时间: {stats.get('reading_time_minutes', '-')} 分钟"
                ]
                sections['conclusion'] = '\n'.join(f"- {x}" for x in st_lines)
            
            # Ensure summary.zh for PDF/Word exporters
            if 'summary' not in report_data or not isinstance(report_data['summary'], dict):
                report_data['summary'] = {}
            if 'zh' not in report_data['summary']:
                report_data['summary']['zh'] = '\n'.join(key_highlights) if key_highlights else ''
            
            report_data['sections'] = sections
        # Export to requested format
        exporter = ReportExporter()
        
        # Use report title as filename, fallback to city_industry format
        safe_city = report_data['city'].replace('/', '_').replace('\\', '_') if report_data.get('city') else 'unknown_city'
        safe_industry = report_data['industry'].replace('/', '_').replace('\\', '_') if report_data.get('industry') else 'unknown_industry'
        
        if report_data.get('title'):
            # Clean up title for filename use
            import re
            filename = report_data['title'].replace(' - AIPE区域产业分析小工作台', '')
            # Remove invalid filename characters
            filename = re.sub(r'[^\w\s\u4e00-\u9fff-]', '', filename)
            # Limit length
            if len(filename) > 50:
                filename = filename[:50]
            filename = filename.strip() or f'{safe_city}_{safe_industry}_报告'
        else:
            filename = f"{safe_city}_{safe_industry}_报告_{now_beijing().strftime("%Y%m%d")}"
        
        if format.lower() == 'pdf':
            export_path = exporter.export_to_pdf(report_data, filename)
            return send_file(export_path, as_attachment=True, download_name=f"{filename}.pdf")
        elif format.lower() in ['word', 'docx']:
            export_path = exporter.export_to_word(report_data, filename)
            return send_file(export_path, as_attachment=True, download_name=f"{filename}.docx")
        elif format.lower() in ['excel', 'xlsx']:
            export_path = exporter.export_to_excel(report_data, filename)
            return send_file(export_path, as_attachment=True, download_name=f"{filename}.xlsx")
        else:
            return jsonify({'error': '不支持的导出格式'}), 400
    
    except Exception as e:
        logger.error(f"Error exporting report: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/reports/<report_id>/download-json', methods=['GET'])
@login_required
def api_download_report_json(report_id):
    try:
        from flask import send_file
        report = Report.query.filter_by(report_id=report_id).first()
        if not report or (report.user_id != current_user.id and current_user.role != 'admin'):
            return jsonify({'error': '无权限访问'}), 403
        file_path = _resolve_report_file_path(report_id, report)
        if not file_path:
            return jsonify({'error': '报告文件不存在'}), 404
        safe_city = (report.city or 'unknown_city').replace('/', '_').replace('\\', '_')
        safe_industry = (report.industry or 'unknown_industry').replace('/', '_').replace('\\', '_')
        filename = f"{safe_city}_{safe_industry}_产业分析报告.json"
        return send_file(str(file_path), as_attachment=True, download_name=filename)
    except Exception as e:
        logger.error(f"Error downloading report JSON: {e}")
        return jsonify({'error': f'下载失败: {str(e)}'}), 500

# Web scraping API
@app.route('/api/scrape-data', methods=['POST'])
@login_required
def api_scrape_data():
    try:
        data = request.get_json() or {}
        city = data.get('city')
        industry = data.get('industry')
        if not city or not industry:
            return jsonify({'error': '请输入城市和行业名称'}), 400
        from src.data.web_scraper import WebScraper
        scraper = WebScraper()
        results = scraper.scrape_policy_data(city, industry)
        scraper.close()
        return jsonify(results)
    except Exception as e:
        logger.error(f"Error scraping data: {e}")
        return jsonify({'error': str(e)}), 500

# Enterprises API for Baidu map
@app.route('/api/enterprises/<report_id>')
@login_required
def api_enterprises(report_id):
    try:
        report = Report.query.filter_by(report_id=report_id).first()
        if not report or (report.user_id != current_user.id and current_user.role != 'admin'):
            return jsonify({'error': '无权限访问'}), 403
        from src.data.web_scraper import WebScraper
        scraper = WebScraper()
        enterprises = scraper.scrape_enterprise_data(report.city or '', report.industry or '', limit=50)
        scraper.close()
        return jsonify({'enterprises': enterprises, 'city': report.city, 'industry': report.industry})
    except Exception as e:
        logger.error(f"Error getting enterprises: {e}")
        return jsonify({'error': str(e)}), 500

# Baidu map page
@app.route('/baidu-map/<report_id>')
@login_required
def baidu_map_page(report_id):
    try:
        report = Report.query.filter_by(report_id=report_id).first()
        if not report or (report.user_id != current_user.id and current_user.role != 'admin'):
            flash('无权限访问该报告')
            return redirect(url_for('index'))
        # Load config for AK
        import json
        with open('config.json', 'r', encoding='utf-8') as f:
            cfg = json.load(f)
        ak = cfg.get('baidu_map_ak', '')
        return render_template('baidu_map.html', report=report, ak=ak)
    except Exception as e:
        logger.error(f"Error rendering baidu map page: {e}")
        flash('加载百度地图失败')
        return redirect(url_for('index'))

# Provenance API (溯源)
@app.route('/api/report/<report_id>/source')
@login_required
def api_report_source(report_id):
    try:
        query = request.args.get('query', '')
        report = Report.query.filter_by(report_id=report_id).first()
        if not report or (report.user_id != current_user.id and current_user.role != 'admin'):
            return jsonify({'error': '无权限访问'}), 403
        if report.report_type == 'llm':
            default_path = APP_ROOT / 'data' / 'output' / 'llm_reports' / f"{report_id}.json"
            file_path = default_path if default_path.exists() else Path(report.file_path)
        else:
            file_path = Path(report.file_path)
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        text = data.get('full_content', '')
        # Return sentences containing query
        import re
        sentences = re.split(r'[。.!?\n]', text)
        matches = [s.strip() for s in sentences if query and query in s][:5]
        return jsonify({'query': query, 'matches': matches})
    except Exception as e:
        logger.error(f"Error in provenance: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/config', methods=['GET', 'POST'])
@login_required
def api_config():
    """API endpoint for configuration management."""
    config_path = 'config.json'
    
    if request.method == 'GET':
        try:
            if os.path.exists(config_path):
                with open(config_path, 'r', encoding='utf-8') as f:
                    config = json.load(f)
            else:
                config = get_default_config()
            return jsonify(config)
        except Exception as e:
            logger.error(f"Error reading config: {e}")
            return jsonify({'error': '读取配置失败'}), 500
    
    elif request.method == 'POST':
        try:
            new_cfg = request.json or {}
            existing = {}
            if os.path.exists(config_path):
                try:
                    with open(config_path, 'r', encoding='utf-8') as f:
                        existing = json.load(f)
                except Exception:
                    existing = {}
            # Preserve api_keys and merge others
            merged = existing.copy()
            # Deep merge api_keys
            merged_api = (existing.get('api_keys') or {}).copy()
            for k, v in (new_cfg.get('api_keys') or {}).items():
                merged_api[k] = v
            if merged_api:
                merged['api_keys'] = merged_api
            # Merge top-level keys
            for k, v in new_cfg.items():
                if k == 'api_keys':
                    continue
                merged[k] = v
            # Add/Update version field
            merged['version'] = '1.1'
            # Write back with indent=2
            with open(config_path, 'w', encoding='utf-8') as f:
                json.dump(merged, f, ensure_ascii=False, indent=2)
            return jsonify({'success': True, 'version': merged['version']})
        except Exception as e:
            logger.error(f"Error saving config: {e}")
            return jsonify({'error': '保存配置失败'}), 500


def get_default_config():
    """Get default configuration."""
    return {
        "categories": [
            "产业概述", "政策环境", "市场规模", "重点企业",
            "技术趋势", "发展机遇", "挑战风险", "未来展望"
        ],
        "ai_integration_focus": [
            "智能制造", "数据分析", "自动化流程", "预测性维护",
            "供应链优化", "客户服务", "质量控制"
        ]
    }


@app.route('/poi-map-visualization')
@login_required
def poi_map_visualization():
    """POI Map Visualization page for regional industry analysis."""
    # Load config
    config_path = 'config.json'
    config = {}
    if os.path.exists(config_path):
        with open(config_path, 'r', encoding='utf-8') as f:
            config = json.load(f)

    # Get API keys
    baidu_ak = config.get('api_keys', {}).get('baidu_map', '')
    google_map_key = config.get('api_keys', {}).get('google_map', '')
    google_map_id = (config.get('google_maps', {}) or {}).get('map_id', '')

    return render_template('poi_map_visualization.html',
                           config=config,
                           google_map_key=google_map_key,
                           google_map_id=google_map_id,
                           baidu_ak=baidu_ak)

@app.route('/api/boundaries', methods=['GET'])
@login_required
def api_boundaries():
    try:
        region = (request.args.get('region') or '').strip()
        level = (request.args.get('level') or 'city').strip()
        bbox = request.args.get('bbox')
        if not region:
            return jsonify({'error': '缺少region参数'}), 400
        mapping_path = APP_ROOT / 'data' / 'boundaries' / 'mapping.json'
        mapping = {}
        if mapping_path.exists():
            with open(mapping_path, 'r', encoding='utf-8') as f:
                mapping = json.load(f)
        regions = mapping.get('regions', {})
        entry = regions.get(region)
        if not entry or not isinstance(entry.get('levels', {}), dict) or not entry['levels'].get(level):
            return jsonify({'error': '未找到该区域的精确边界数据映射', 'region': region, 'level': level}), 404
        src = entry['levels'][level]
        if isinstance(src, str) and src.startswith(('http://', 'https://')):
            import requests
            r = requests.get(src, timeout=8)
            if r.status_code != 200:
                return jsonify({'error': f'边界数据源不可用: HTTP {r.status_code}'}), 502
            gj = r.json()
        else:
            p = APP_ROOT / src
            if not p.exists():
                return jsonify({'error': f'本地边界文件不存在: {str(p)}'}), 404
            with open(p, 'r', encoding='utf-8') as f:
                gj = json.load(f)
        if bbox:
            try:
                minLng, minLat, maxLng, maxLat = [float(x) for x in bbox.split(',')]
                def geom_in_bbox(geom):
                    try:
                        coords_iter = []
                        t = geom.get('type')
                        c = geom.get('coordinates')
                        if t == 'Polygon':
                            for ring in c:
                                coords_iter.extend(ring)
                        elif t == 'MultiPolygon':
                            for poly in c:
                                for ring in poly:
                                    coords_iter.extend(ring)
                        elif t == 'LineString':
                            coords_iter.extend(c)
                        elif t == 'MultiLineString':
                            for line in c:
                                coords_iter.extend(line)
                        elif t == 'Point':
                            coords_iter.append(c)
                        elif t == 'MultiPoint':
                            coords_iter.extend(c)
                        else:
                            return True
                        for lng, lat in coords_iter:
                            if (minLng <= lng <= maxLng) and (minLat <= lat <= maxLat):
                                return True
                        return False
                    except Exception:
                        return True
                if gj.get('type') == 'FeatureCollection':
                    gj['features'] = [f for f in gj.get('features', []) if geom_in_bbox(f.get('geometry', {}))]
            except Exception:
                pass
        meta = entry.get('meta', {})
        return jsonify({'success': True, 'region': region, 'level': level, 'meta': meta, 'geojson': gj})
    except Exception as e:
        logger.error(f"Error in api_boundaries: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/poi-search', methods=['POST'])
@login_required
def api_poi_search():
    """API endpoint for searching POIs by region and keywords."""
    try:
        data = request.get_json()
        region = data.get('region', '').strip()
        keywords = data.get('keywords', [])

        if not region:
            return jsonify({'error': '请输入目标区域'}), 400

        if not keywords or not isinstance(keywords, list):
            return jsonify({'error': '请提供至少一个搜索标签'}), 400

        # Load Google Map API key from config (prioritize Google)
        config_path = 'config.json'
        google_map_key = ''
        if os.path.exists(config_path):
            with open(config_path, 'r', encoding='utf-8') as f:
                config = json.load(f)
                google_map_key = config.get('api_keys', {}).get('google_map', '')

        if not google_map_key:
            return jsonify({'error': 'Google地图API密钥未配置'}), 500

        # Perform actual POI search using Google Places API
        from src.analysis.poi_searcher import GooglePlacesSearcher, PoiDataProcessor, PoiExporter

        searcher = GooglePlacesSearcher(google_map_key)
        pois = searcher.batch_search_poi(region, keywords, max_results=100)

        if not pois:
            return jsonify({
                'success': True,
                'data': {
                    'results': [],
                    'total_count': 0,
                    'page_info': {'current': 1, 'size': 0, 'total': 1},
                    'export_paths': {
                        'json': '',
                        'excel': ''
                    },
                    'visualization': {
                        'heatmap_data': [],
                        'cluster_data': []
                    }
                }
            })

        # Process the data for visualization
        processor = PoiDataProcessor()
        heatmap_data = processor.generate_heatmap_data(pois)
        cluster_data = processor.generate_cluster_data(pois)

        # Generate export paths
        timestamp = now_beijing().strftime('%Y%m%d_%H%M%S')
        search_id = f"poi_{timestamp}"
        base_path = Path('data/output/poi_search') / search_id
        base_path.mkdir(parents=True, exist_ok=True)

        json_path = str(base_path / f"{search_id}.json")
        excel_path = str(base_path / f"{search_id}.xlsx")

        # Export data
        exporter = PoiExporter()
        exporter.export_to_json(pois, json_path)
        exporter.export_to_excel(pois, excel_path)

        # Calculate statistics
        stats = processor.calculate_statistics(pois)

        return jsonify({
            'success': True,
            'data': {
                'results': pois,
                'total_count': len(pois),
                'page_info': {'current': 1, 'size': len(pois), 'total': 1},
                'export_paths': {
                    'json': json_path,
                    'excel': excel_path
                },
                'visualization': {
                    'heatmap_data': heatmap_data,
                    'cluster_data': cluster_data
                },
                'statistics': stats
            }
        })

    except Exception as e:
        logger.error(f"Error in POI search API: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return jsonify({'error': f'POI搜索失败: {str(e)}'}), 500


@app.route('/api/poi-upload', methods=['POST'])
@login_required
def api_poi_upload():
    """API endpoint for uploading and parsing POI data files."""
    try:
        if 'file' not in request.files:
            return jsonify({'error': '没有上传文件'}), 400

        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': '没有选择文件'}), 400

        if file:
            # Save uploaded file temporarily
            import tempfile
            import os
            from werkzeug.utils import secure_filename
            from src.analysis.poi_parser import PoiDocumentParser
            from src.analysis.poi_searcher import PoiDataProcessor, PoiExporter

            # Create temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(file.filename)[1]) as tmp_file:
                file.save(tmp_file.name)
                temp_path = tmp_file.name

            try:
                # Parse the file
                parser = PoiDocumentParser()
                pois = parser.parse_file(temp_path)

                # Validate and normalize coordinates using pre-stored location fields
                valid_pois = []
                invalid_count = 0
                invalid_breakdown = {
                    'missing_or_unparsed': 0,
                    'out_of_range': 0
                }
                invalid_examples = []
                for poi in pois:
                    loc = poi.get('location') or {}
                    # Support both {lng, lat} and {longitude, latitude}
                    lng = loc.get('lng', loc.get('longitude'))
                    lat = loc.get('lat', loc.get('latitude'))
                    try:
                        if lng is not None and lat is not None:
                            lng = float(lng)
                            lat = float(lat)
                            if -180.0 <= lng <= 180.0 and -90.0 <= lat <= 90.0:
                                poi.setdefault('location', {})
                                poi['location']['lng'] = lng
                                poi['location']['lat'] = lat
                                valid_pois.append(poi)
                            else:
                                invalid_count += 1
                                invalid_breakdown['out_of_range'] += 1
                                if len(invalid_examples) < 5:
                                    invalid_examples.append({'name': poi.get('name', ''), 'reason': '坐标超出范围', 'lng': lng, 'lat': lat})
                        else:
                            invalid_count += 1
                            invalid_breakdown['missing_or_unparsed'] += 1
                            if len(invalid_examples) < 5:
                                invalid_examples.append({'name': poi.get('name', ''), 'reason': '缺少或无法解析坐标'})
                    except Exception:
                        invalid_count += 1
                        invalid_breakdown['missing_or_unparsed'] += 1
                        if len(invalid_examples) < 5:
                            invalid_examples.append({'name': poi.get('name', ''), 'reason': '坐标解析异常'})

                # Generate export paths
                timestamp = now_beijing().strftime('%Y%m%d_%H%M%S')
                upload_id = f"poi_upload_{timestamp}"
                base_path = Path('data/output/poi_upload') / upload_id
                base_path.mkdir(parents=True, exist_ok=True)

                json_path = str(base_path / f"{upload_id}.json")
                excel_path = str(base_path / f"{upload_id}.xlsx")

                # Export data using the exporter
                exporter = PoiExporter()
                exporter.export_to_json(valid_pois, json_path)
                exporter.export_to_excel(valid_pois, excel_path)

                # Process data for visualization
                processor = PoiDataProcessor()
                heatmap_data = processor.generate_heatmap_data(valid_pois)
                cluster_data = processor.generate_cluster_data(valid_pois)
                stats = processor.calculate_statistics(valid_pois)
                stats.update({
                    'parse_processed': len(pois),
                    'parse_valid': len(valid_pois),
                    'parse_invalid': invalid_count,
                    'invalid_breakdown': invalid_breakdown
                })

                # Remove temporary file
                os.unlink(temp_path)

                return jsonify({
                    'success': True,
                    'data': {
                        'results': valid_pois,
                        'total_count': len(valid_pois),
                        'parse': {
                            'processed': len(pois),
                            'valid': len(valid_pois),
                            'invalid': invalid_count,
                            'invalid_breakdown': invalid_breakdown,
                            'invalid_examples': invalid_examples
                        },
                        'export_paths': {
                            'json': json_path,
                            'excel': excel_path
                        },
                        'visualization': {
                            'heatmap_data': heatmap_data,
                            'cluster_data': cluster_data
                        },
                        'statistics': stats
                    }
                })

            except Exception as e:
                # Remove temporary file in case of error
                if os.path.exists(temp_path):
                    os.unlink(temp_path)
                logger.error(f"Error processing POI file: {e}")
                import traceback
                logger.error(traceback.format_exc())
                return jsonify({'error': f'文件处理失败: {str(e)}'}), 500

    except Exception as e:
        logger.error(f"Error in POI upload API: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return jsonify({'error': f'上传失败: {str(e)}'}), 500


@app.route('/about')
def about_page():
    """Renders the about page by converting Project Summary.md to HTML."""
    try:
        summary_path = APP_ROOT / 'Project Summary.md'
        if summary_path.exists():
            with open(summary_path, 'r', encoding='utf-8') as f:
                markdown_content = f.read()
            
            html_content = markdown2.markdown(
                markdown_content,
                extras=[
                    "fenced-code-blocks",
                    "tables",
                    "strike",
                    "task_list",
                    "code-friendly",
                    "toc",
                ],
            )
        else:
            html_content = "<h1>关于文件未找到</h1><p>Project Summary.md 文件不存在。</p>"
    except Exception as e:
        logger.error(f"Error reading or converting Project Summary.md: {e}")
        html_content = "<h1>错误</h1><p>加载关于页面时出错。</p>"
        
    return render_template('about.html', content=html_content)


def main():
    """Main function to run the app."""
    # Initialize database
    init_db()

    # Setup directories
    setup_directories()

    # Register notification routes
    register_notification_routes(app)

    # Register streaming blueprint
    app.register_blueprint(streaming_bp, url_prefix='/streaming')
    app.register_blueprint(report_gen_bp)
    app.register_blueprint(report_generation_bp)

    # Fetch WeChat articles in the background when the app starts
    try:
        from src.tasks.wechat_tasks import fetch_wechat_articles_task
        logger.info("Starting background WeChat articles fetch task...")
        fetch_wechat_articles_task.delay()  # Run asynchronously
        logger.info("Background WeChat articles fetch task started")
    except Exception as e:
        logger.error(f"Error starting WeChat articles fetch task: {e}")

    # Run Flask app
    # In a production environment, use a proper WSGI server like Gunicorn or uWSGI
    # Example: gunicorn --workers 4 --bind 0.0.0.0:8000 app:app
    app.run(debug=True, port=5000, host='0.0.0.0')



@app.route('/api/wechat-config', methods=['GET'])
def get_wechat_config():
    try:
        with open('data/wechat_accounts_config.json', 'r', encoding='utf-8') as f:
            config = json.load(f)
        return jsonify(config)
    except FileNotFoundError:
        return jsonify([])
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/wechat-config', methods=['POST'])
def save_wechat_config():
    try:
        config_data = request.json
        with open('data/wechat_accounts_config.json', 'w', encoding='utf-8') as f:
            json.dump(config_data, f, ensure_ascii=False, indent=2)
        return jsonify({"status": "success"})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


if __name__ == '__main__':
    # Refresh leadership data on startup
    print("Refreshing leadership data on startup...")
    regions_to_refresh = ['chengdu', 'sichuan', 'beijing', 'shanghai', 'chongqing', 'shenzhen', 'wuhan']
    for region in regions_to_refresh:
        try:
            leadership_data = leadership_scraper.get_leadership_data(region)
            print(f"✓ Refreshed leadership data for {region}")
        except Exception as e:
            print(f"✗ Error refreshing {region}: {e}")

    print("Leadership data refresh completed.")

    # webbrowser.open('http://127.0.0.1:5000')
    main()
@app.route('/home/a')
@login_required
def home_layout_a():
    return render_template('home_layout_a.html')

@app.route('/home/b')
@login_required
def home_layout_b():
    return render_template('home_layout_b.html')

@app.route('/home/c')
@login_required
def home_layout_c():
    return render_template('home_layout_c.html')


# Mindmap API
@app.route('/api/policy-mindmap', methods=['POST'])
@login_required
def api_policy_mindmap():
    """Generate mindmap from policy document using ERNIE Bot enhanced parsing"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': '未上传文件'}), 400
        file = request.files['file']
        filename = file.filename
        ext = filename.rsplit('.', 1)[-1].lower()
        content = ''

        # File parsing (similar to existing parsing)
        import tempfile
        import os

        with tempfile.NamedTemporaryFile(delete=False, suffix=f'.{ext}') as temp_file:
            file.save(temp_file.name)
            temp_path = temp_file.name

        try:
            # File parsing based on format
            if ext in ['txt', 'md']:
                with open(temp_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
            elif ext in ['doc', 'docx']:
                if ext == 'docx':
                    from docx import Document
                    doc = Document(temp_path)
                    content = '\n'.join(p.text for p in doc.paragraphs)
                else:
                    try:
                        import docx2txt
                        content = docx2txt.process(temp_path)
                    except ImportError:
                        try:
                            from docx import Document
                            doc = Document(temp_path)
                            content = '\n'.join(p.text for p in doc.paragraphs)
                        except Exception:
                            content = "DOC文件解析失败，请尝试转换为DOCX格式后重新上传。"
            elif ext == 'pdf':
                try:
                    import PyPDF2
                    with open(temp_path, 'rb') as f:
                        reader = PyPDF2.PdfReader(f)
                        for page in reader.pages:
                            content += page.extract_text() or ''
                except Exception:
                    try:
                        import pypdf
                        with open(temp_path, 'rb') as f:
                            reader = pypdf.PdfReader(f)
                            for page in reader.pages:
                                content += page.extract_text() or ''
                    except Exception:
                        content = "PDF文件解析失败，请检查文件格式是否正确。"
            else:
                return jsonify({'error': '不支持的文件格式'}), 400
        finally:
            os.unlink(temp_path)

        # Use hybrid parser to get structured content
        from hybrid_parser import HybridPolicyParser
        parser = HybridPolicyParser()
        structured_content = parser.parse_policy_document(content, llm_service='kimi')  # Use Kimi as default for mindmap

        # Generate mindmap data from structured content
        from mindmap_generator import MindMapGenerator
        mindmap_gen = MindMapGenerator()
        mindmap_data = mindmap_gen.generate_mindmap_data(structured_content)

        return jsonify({
            'success': True,
            'mindmap': mindmap_data
        })
    except Exception as e:
        logger.error(f"Error in mindmap generation: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return jsonify({'error': str(e)}), 500


# API endpoint to fetch regional data including leadership info
@app.route("/api/regional-data/<region>", methods=["GET"])
@login_required
def api_regional_data(region):
    """Fetch real-time regional data including leadership info"""
    try:
        from src.data.regional_data_scraper import RegionalDataScraper
        scraper = RegionalDataScraper()

        # Fetch leadership and economic data
        leadership_data = scraper.get_regional_leadership(region)
        gdp_data = scraper.get_regional_gdp_data(region)
        district_rankings = scraper.get_district_rankings(region)
        university_info = scraper.get_university_info(region)
        sci_institutions = scraper.get_science_institutions(region)

        return jsonify({
            "success": True,
            "region": region,
            "leadership": leadership_data,
            "gdp_data": gdp_data,
            "district_rankings": district_rankings,
            "universities": university_info,
            "research_institutions": sci_institutions
        })
    except Exception as e:
        logger.error(f"Error fetching regional data: {e}")
        return jsonify({"error": str(e)}), 500


# Route to serve the local POI data file for demo
@app.route('/data/output/poi_upload/chengdu-ai-pois.json')
def serve_local_poi_data():
    """Serve the local POI data file for the demo."""
    from flask import send_from_directory
    import os

    # Ensure we use the correct absolute path relative to the app root
    base_dir = os.path.dirname(os.path.abspath(__file__))  # Get the directory of the current script
    file_path = os.path.join(base_dir, 'data', 'output', 'poi_upload')
    filename = 'chengdu-ai-pois.json'

    return send_from_directory(file_path, filename)
