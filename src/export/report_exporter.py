#!/usr/bin/env python3
"""
Report Export Module
Export reports to various formats: PDF, Word, Excel
"""

import os
import json
import logging
from datetime import datetime
from pathlib import Path
from typing import Dict, Optional

# PDF Generation
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.enums import TA_CENTER, TA_LEFT

# Word Generation
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Excel Generation
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill
from openpyxl.chart import BarChart, Reference

logger = logging.getLogger(__name__)


class ReportExporter:
    """Export reports to various formats."""
    
    def __init__(self, output_dir='data/output/exports'):
        """Initialize the report exporter.
        
        Args:
            output_dir: Directory to save exported reports
        """
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self._register_chinese_fonts()
    
    def _register_chinese_fonts(self):
        """Register Chinese fonts for PDF generation."""
        try:
            font_paths = [
                '/System/Library/Fonts/STHeiti Medium.ttc',
                '/System/Library/Fonts/Supplemental/Songti.ttc',
                '/System/Library/Fonts/PingFang.ttc'
            ]
            
            for font_path in font_paths:
                if os.path.exists(font_path):
                    try:
                        pdfmetrics.registerFont(TTFont('Chinese', font_path))
                        logger.info(f"Registered Chinese font: {font_path}")
                        return
                    except Exception as e:
                        logger.warning(f"Failed to register font {font_path}: {e}")
                        continue
            
            logger.warning("No Chinese fonts found, PDF may not display Chinese correctly")
        except Exception as e:
            logger.error(f"Error registering Chinese fonts: {e}")
    
    def _clean_content_formatting(self, content: str) -> str:
        """Remove markdown and JSON formatting from content."""
        import re
        
        # Remove JSON code block markers
        content = re.sub(r'^```json\s*', '', content)
        content = re.sub(r'^```\s*', '', content)
        content = re.sub(r'\s*```$', '', content)
        
        # Remove markdown code block markers
        content = re.sub(r'```.*?\n', '', content)
        content = re.sub(r'```\s*$', '', content)
        
        # Remove markdown headers
        content = re.sub(r'^#+\s*', '', content)
        
        # Convert markdown bold to HTML bold
        content = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', content)
        content = re.sub(r'__(.*?)__', r'<b>\1</b>', content)
        
        # Convert markdown italic to HTML italic
        content = re.sub(r'\*(.*?)\*', r'<i>\1</i>', content)
        content = re.sub(r'_(.*?)_', r'<i>\1</i>', content)
        
        # Convert markdown line breaks to HTML line breaks
        content = content.replace('\n', '<br/>')
        
        # Clean up extra whitespace
        content = content.strip()
        
        return content
    
    def export_to_pdf(self, report_data: Dict, filename: str) -> str:
        """Export report to PDF format.
        
        Args:
            report_data: Report data dictionary
            filename: Output filename (without extension)
        
        Returns:
            Path to the generated PDF file
        """
        try:
            output_path = self.output_dir / f"{filename}.pdf"
            
            # Create PDF document
            doc = SimpleDocTemplate(
                str(output_path),
                pagesize=A4,
                topMargin=0.75*inch,
                bottomMargin=0.75*inch,
                leftMargin=1*inch,
                rightMargin=1*inch
            )
            
            # Container for PDF elements
            story = []
            styles = getSampleStyleSheet()
            
            # Custom styles with Chinese font support
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Title'],
                fontSize=24,
                textColor=colors.HexColor('#667eea'),
                spaceAfter=30,
                alignment=TA_CENTER,
                fontName='Chinese'
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading1'],
                fontSize=16,
                textColor=colors.HexColor('#667eea'),
                spaceAfter=12,
                spaceBefore=12,
                fontName='Chinese'
            )
            
            normal_style = ParagraphStyle(
                'ChineseNormal',
                parent=styles['Normal'],
                fontSize=11,
                fontName='Chinese',
                leading=16
            )
            
            # Improved styles for better appearance
            subtitle_style = ParagraphStyle(
                'Subtitle',
                parent=heading_style,
                fontSize=14,
                textColor=colors.HexColor('#7c3aed'),
                spaceAfter=10,
                spaceBefore=10,
                fontName='Chinese'
            )
            
            section_style = ParagraphStyle(
                'Section',
                parent=normal_style,
                fontSize=12,
                leading=18,
                fontName='Chinese'
            )
            
            # Add title
            title = report_data.get('title', f"{report_data.get('city', 'N/A')} {report_data.get('industry', 'N/A')} 产业分析报告")
            # Clean up title if it's too long
            if 'AIPE区域产业分析小工作台' in title:
                title = title.replace(' - AIPE区域产业分析小工作台', '')
            # Remove markdown and JSON formatting from title
            title = self._clean_content_formatting(title)
            story.append(Paragraph(title, title_style))
            story.append(Spacer(1, 0.3*inch))
            
            # Add metadata
            if report_data.get('metadata') or report_data.get('generated_at'):
                meta_text = f"生成时间: {report_data.get('generated_at', report_data.get('metadata', {}).get('processed_at', 'N/A'))}"
                # Remove markdown and JSON formatting from metadata
                meta_text = self._clean_content_formatting(meta_text)
                story.append(Paragraph(meta_text, normal_style))
                story.append(Spacer(1, 0.2*inch))
            
            # Add Executive Summary from dashboard data
            if report_data.get('summary'):
                story.append(Paragraph("执行摘要", heading_style))
                summary = report_data['summary']
                if isinstance(summary, dict):
                    # Add key metrics
                    metrics = []
                    if summary.get('word_count'):
                        metrics.append(f"总字数: {summary['word_count']}")
                    if summary.get('reading_time'):
                        metrics.append(f"阅读时间: {summary['reading_time']} 分钟")
                    if summary.get('categories_analyzed'):
                        metrics.append(f"分析类别: {summary['categories_analyzed']}")
                    if summary.get('ai_opportunities'):
                        metrics.append(f"AI机会: {summary['ai_opportunities']}")
                    
                    if metrics:
                        metrics_text = " | ".join(metrics)
                        story.append(Paragraph(metrics_text, section_style))
                        story.append(Spacer(1, 0.1*inch))
                    
                    # Add key highlights
                    if summary.get('key_highlights'):
                        story.append(Paragraph("关键发现", subtitle_style))
                        for highlight in summary['key_highlights']:
                            story.append(Paragraph(f"• {highlight}", section_style))
                            story.append(Spacer(1, 0.05*inch))
                story.append(Spacer(1, 0.2*inch))
            
            # Add SWOT Analysis
            if report_data.get('swot_analysis'):
                swot = report_data['swot_analysis']
                if any([swot.get('strengths'), swot.get('weaknesses'), 
                       swot.get('opportunities'), swot.get('threats')]):
                    story.append(PageBreak())
                    story.append(Paragraph("SWOT 战略分析", heading_style))
                    
                    swot_sections = [
                        ('优势 (Strengths)', swot.get('strengths', []), colors.HexColor('#10b981')),
                        ('劣势 (Weaknesses)', swot.get('weaknesses', []), colors.HexColor('#ef4444')),
                        ('机遇 (Opportunities)', swot.get('opportunities', []), colors.HexColor('#3b82f6')),
                        ('威胁 (Threats)', swot.get('threats', []), colors.HexColor('#f59e0b'))
                    ]
                    
                    for section_title, items, color in swot_sections:
                        if items:
                            section_style = ParagraphStyle(
                                'SectionTitle',
                                parent=heading_style,
                                fontSize=14,
                                textColor=color,
                                fontName='Chinese'
                            )
                            story.append(Paragraph(section_title, section_style))
                            for i, item in enumerate(items, 1):
                                # Remove markdown and JSON formatting
                                clean_item = self._clean_content_formatting(item)
                                item_text = f"{i}. {clean_item}"
                                story.append(Paragraph(item_text, normal_style))
                                story.append(Spacer(1, 0.05*inch))
                            story.append(Spacer(1, 0.15*inch))
            
            # Add sections or full content
            if report_data.get('sections'):
                for section_key, section_content in report_data['sections'].items():
                    story.append(PageBreak())
                    section_titles = {
                        'executive_summary': '执行摘要',
                        'industry_overview': '产业概览',
                        'policy_landscape': '政策环境',
                        'ecosystem': '产业生态',
                        'value_chain': '产业链分析',
                        'ai_integration': 'AI融合潜力',
                        'conclusion': '结论与建议'
                    }
                    story.append(Paragraph(section_titles.get(section_key, section_key), heading_style))
                    
                    for para in section_content.split('\n\n'):
                        if para.strip():
                            # Remove markdown and JSON formatting
                            clean_para = self._clean_content_formatting(para.strip())
                            if len(clean_para) > 10:
                                story.append(Paragraph(clean_para, normal_style))
                                story.append(Spacer(1, 0.1*inch))
            elif report_data.get('full_content'):
                story.append(PageBreak())
                story.append(Paragraph("完整报告", heading_style))
                full_content = report_data.get('full_content', '')
                for para in full_content.split('\n\n'):
                    if para.strip() and not para.startswith('#'):
                        clean_para = self._clean_content_formatting(para.strip().replace('#', ''))
                        if len(clean_para) > 10:
                            story.append(Paragraph(clean_para, normal_style))
                            story.append(Spacer(1, 0.1*inch))
            else:
                # Handle uploaded analysis reports (dashboard data)
                story.append(PageBreak())
                story.append(Paragraph("分析报告", heading_style))
                
                # Add summary
                if report_data.get('summary'):
                    story.append(Paragraph("报告摘要", heading_style))
                    summary = report_data['summary']
                    if isinstance(summary, dict):
                        # Add key metrics
                        metrics_text = f"总字数: {summary.get('word_count', 0)} | 阅读时间: {summary.get('reading_time', 0)} 分钟 | 分析类别: {summary.get('categories_analyzed', 0)}"
                        story.append(Paragraph(metrics_text, normal_style))
                        story.append(Spacer(1, 0.1*inch))
                        
                        # Add key highlights
                        if summary.get('key_highlights'):
                            story.append(Paragraph("关键发现", heading_style))
                            for highlight in summary['key_highlights']:
                                story.append(Paragraph(f"• {highlight}", normal_style))
                                story.append(Spacer(1, 0.05*inch))
                
                # Add categories analysis
                if report_data.get('categories'):
                    story.append(Paragraph("内容分类分析", heading_style))
                    for category_name, category_data in report_data['categories'].items():
                        if category_data.get('has_content'):
                            story.append(Paragraph(category_name, subtitle_style))
                            # Add relevance score
                            if category_data.get('relevance_score'):
                                score_text = f"相关性得分: {category_data['relevance_score']}"
                                story.append(Paragraph(score_text, section_style))
                            # Add key points
                            if category_data.get('key_points'):
                                for point in category_data['key_points']:
                                    story.append(Paragraph(f"• {point}", section_style))
                                    story.append(Spacer(1, 0.05*inch))
                            story.append(Spacer(1, 0.1*inch))
                
                # Add AI opportunities with scores
                if report_data.get('ai_opportunities'):
                    story.append(Paragraph("AI应用机会", heading_style))
                    # Create a table for AI opportunities
                    ai_data = []
                    ai_data.append(['应用领域', '潜力得分', '优先级', '建议'])
                    for ai_name, ai_data_item in report_data['ai_opportunities'].items():
                        score = ai_data_item.get('potential_score', 0)
                        priority = ai_data_item.get('priority_level', 'none')
                        priority_text = {'high': '高', 'medium': '中', 'low': '低', 'none': '无'}.get(priority, priority)
                        recommendation = ai_data_item.get('recommendation', 'N/A')
                        ai_data.append([ai_name, str(score), priority_text, recommendation])
                    
                    # Create table
                    ai_table = Table(ai_data, colWidths=[1.5*inch, 0.8*inch, 0.6*inch, 2.5*inch])
                    ai_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#667eea')),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Chinese'),
                        ('FONTSIZE', (0, 0), (-1, 0), 10),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                        ('FONTNAME', (0, 1), (-1, -1), 'Chinese'),
                        ('FONTSIZE', (0, 1), (-1, -1), 9),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black)
                    ]))
                    story.append(ai_table)
                    story.append(Spacer(1, 0.2*inch))
                
                # Add key insights
                if report_data.get('key_insights'):
                    story.append(Paragraph("关键洞察", heading_style))
                    for insight in report_data['key_insights']:
                        if isinstance(insight, dict):
                            insight_title = insight.get('title', '洞察')
                            story.append(Paragraph(insight_title, subtitle_style))
                            insight_text = insight.get('text') or str(insight.get('data', ''))
                            if insight_text:
                                story.append(Paragraph(insight_text, section_style))
                        else:
                            story.append(Paragraph(str(insight), section_style))
                        story.append(Spacer(1, 0.1*inch))
                
                # Add charts data as tables
                if report_data.get('charts'):
                    story.append(PageBreak())
                    story.append(Paragraph("数据分析图表", heading_style))
                    
                    charts = report_data['charts']
                    
                    # Category distribution chart data
                    if charts.get('category_distribution'):
                        story.append(Paragraph("内容分布分析", subtitle_style))
                        cat_data = charts['category_distribution']
                        if cat_data.get('data'):
                            table_data = [['类别', '得分']]
                            for i, label in enumerate(cat_data['data'][0].get('labels', [])):
                                value = cat_data['data'][0].get('chart_values', [])[i] if i < len(cat_data['data'][0].get('chart_values', [])) else 0
                                table_data.append([label, str(value)])
                            
                            table = Table(table_data, colWidths=[2*inch, 1*inch])
                            table.setStyle(TableStyle([
                                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#667eea')),
                                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                                ('FONTNAME', (0, 0), (-1, 0), 'Chinese'),
                                ('FONTSIZE', (0, 0), (-1, 0), 10),
                                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                                ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                                ('FONTNAME', (0, 1), (-1, -1), 'Chinese'),
                                ('FONTSIZE', (0, 1), (-1, -1), 9),
                                ('GRID', (0, 0), (-1, -1), 1, colors.black)
                            ]))
                            story.append(table)
                            story.append(Spacer(1, 0.2*inch))
                    
                    # AI opportunities chart data
                    if charts.get('ai_opportunities'):
                        story.append(Paragraph("AI应用潜力评估", subtitle_style))
                        ai_chart_data = charts['ai_opportunities']
                        if ai_chart_data.get('data'):
                            table_data = [['应用领域', '潜力得分']]
                            for i, label in enumerate(ai_chart_data['data'][0].get('theta', [])):
                                value = ai_chart_data['data'][0].get('r', [])[i] if i < len(ai_chart_data['data'][0].get('r', [])) else 0
                                table_data.append([label, str(value)])
                            
                            table = Table(table_data, colWidths=[2*inch, 1*inch])
                            table.setStyle(TableStyle([
                                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#7c3aed')),
                                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                                ('FONTNAME', (0, 0), (-1, 0), 'Chinese'),
                                ('FONTSIZE', (0, 0), (-1, 0), 10),
                                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                                ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                                ('FONTNAME', (0, 1), (-1, -1), 'Chinese'),
                                ('FONTSIZE', (0, 1), (-1, -1), 9),
                                ('GRID', (0, 0), (-1, -1), 1, colors.black)
                            ]))
                            story.append(table)
                            story.append(Spacer(1, 0.2*inch))
                    
                    # Keyword frequency chart data
                    if charts.get('keyword_frequency'):
                        story.append(Paragraph("关键词频率分析", subtitle_style))
                        keyword_data = charts['keyword_frequency']
                        if keyword_data.get('data'):
                            table_data = [['关键词', '频率']]
                            for i, keyword in enumerate(keyword_data['data'][0].get('x', [])[:15]):  # Limit to top 15
                                freq = keyword_data['data'][0].get('y', [])[i] if i < len(keyword_data['data'][0].get('y', [])) else 0
                                table_data.append([keyword, str(freq)])
                            
                            table = Table(table_data, colWidths=[2*inch, 1*inch])
                            table.setStyle(TableStyle([
                                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#10b981')),
                                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                                ('FONTNAME', (0, 0), (-1, 0), 'Chinese'),
                                ('FONTSIZE', (0, 0), (-1, 0), 10),
                                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                                ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                                ('FONTNAME', (0, 1), (-1, -1), 'Chinese'),
                                ('FONTSIZE', (0, 1), (-1, -1), 9),
                                ('GRID', (0, 0), (-1, -1), 1, colors.black)
                            ]))
                            story.append(table)
                            story.append(Spacer(1, 0.2*inch))
            
            # Build PDF
            doc.build(story)
            
            logger.info(f"PDF exported successfully: {output_path}")
            return str(output_path)
        
        except Exception as e:
            logger.error(f"Error exporting to PDF: {e}")
            raise
    
    def _clean_content_formatting(self, content: str) -> str:
        """Remove markdown and JSON formatting from content."""
        import re
        
        # Remove JSON code block markers
        content = re.sub(r'^```json\s*', '', content)
        content = re.sub(r'^```\s*', '', content)
        content = re.sub(r'\s*```$', '', content)
        
        # Remove markdown code block markers
        content = re.sub(r'```.*?\n', '', content)
        content = re.sub(r'```\s*$', '', content)
        
        # Remove markdown headers
        content = re.sub(r'^#+\s*', '', content)
        
        # Convert markdown bold to HTML bold
        content = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', content)
        content = re.sub(r'__(.*?)__', r'<b>\1</b>', content)
        
        # Convert markdown italic to HTML italic
        content = re.sub(r'\*(.*?)\*', r'<i>\1</i>', content)
        content = re.sub(r'_(.*?)_', r'<i>\1</i>', content)
        
        # Convert markdown line breaks to HTML line breaks
        content = content.replace('\n', '<br/>')
        
        # Clean up extra whitespace
        content = content.strip()
        
        return content
    
    def export_to_word(self, report_data: Dict, filename: str) -> str:
        """Export report to Word format.
        
        Args:
            report_data: Report data dictionary
            filename: Output filename (without extension)
        
        Returns:
            Path to the generated Word file
        """
        try:
            output_path = self.output_dir / f"{filename}.docx"
            
            # Create Word document
            doc = Document()
            
            # Set document properties
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            # Add title
            title_text = report_data.get('title', f"{report_data.get('city', 'N/A')} {report_data.get('industry', 'N/A')} 产业分析报告")
            # Clean up title if it's too long
            if 'AIPE区域产业分析小工作台' in title_text:
                title_text = title_text.replace(' - AIPE区域产业分析小工作台', '')
            
            title = doc.add_heading(title_text, level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title.style.font.size = Pt(24)
            title.style.font.color.rgb = RGBColor(102, 126, 234)  # #667eea
            
            # Add metadata
            if report_data.get('metadata') or report_data.get('generated_at'):
                meta_text = f"生成时间: {report_data.get('generated_at', report_data.get('metadata', {}).get('processed_at', 'N/A'))}"
                meta = doc.add_paragraph(meta_text)
                meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
                meta.style.font.size = Pt(11)
            doc.add_paragraph()  # Empty line
            
            # Add Executive Summary from dashboard data
            if report_data.get('summary'):
                doc.add_heading('执行摘要', level=1)
                summary = report_data['summary']
                if isinstance(summary, dict):
                    # Add key metrics
                    metrics = []
                    if summary.get('word_count'):
                        metrics.append(f"总字数: {summary['word_count']}")
                    if summary.get('reading_time'):
                        metrics.append(f"阅读时间: {summary['reading_time']} 分钟")
                    if summary.get('categories_analyzed'):
                        metrics.append(f"分析类别: {summary['categories_analyzed']}")
                    if summary.get('ai_opportunities'):
                        metrics.append(f"AI机会: {summary['ai_opportunities']}")
                    
                    if metrics:
                        metrics_text = " | ".join(metrics)
                        metrics_para = doc.add_paragraph(metrics_text)
                        metrics_para.style.font.size = Pt(11)
                        doc.add_paragraph()  # Empty line
                    
                    # Add key highlights
                    if summary.get('key_highlights'):
                        doc.add_heading('关键发现', level=2)
                        for highlight in summary['key_highlights']:
                            p = doc.add_paragraph(highlight, style='List Bullet')
                            p.style.font.size = Pt(11)
                doc.add_paragraph()  # Empty line
            
            # Add SWOT Analysis
            if report_data.get('swot_analysis'):
                swot = report_data['swot_analysis']
                if any([swot.get('strengths'), swot.get('weaknesses'), 
                       swot.get('opportunities'), swot.get('threats')]):
                    doc.add_page_break()
                    doc.add_heading('SWOT 战略分析', level=1)
                    
                    # Strengths
                    if swot.get('strengths'):
                        doc.add_heading('优势 (Strengths)', level=2)
                        for item in swot.get('strengths', []):
                            # Remove markdown and JSON formatting
                            clean_item = self._clean_content_formatting(item)
                            p = doc.add_paragraph(clean_item, style='List Bullet')
                            p.style.font.size = Pt(11)
                    
                    # Weaknesses
                    if swot.get('weaknesses'):
                        doc.add_heading('劣势 (Weaknesses)', level=2)
                        for item in swot.get('weaknesses', []):
                            # Remove markdown and JSON formatting
                            clean_item = self._clean_content_formatting(item)
                            p = doc.add_paragraph(clean_item, style='List Bullet')
                            p.style.font.size = Pt(11)
                    
                    # Opportunities
                    if swot.get('opportunities'):
                        doc.add_heading('机遇 (Opportunities)', level=2)
                        for item in swot.get('opportunities', []):
                            # Remove markdown and JSON formatting
                            clean_item = self._clean_content_formatting(item)
                            p = doc.add_paragraph(clean_item, style='List Bullet')
                            p.style.font.size = Pt(11)
                    
                    # Threats
                    if swot.get('threats'):
                        doc.add_heading('威胁 (Threats)', level=2)
                        for item in swot.get('threats', []):
                            # Remove markdown and JSON formatting
                            clean_item = self._clean_content_formatting(item)
                            p = doc.add_paragraph(clean_item, style='List Bullet')
                            p.style.font.size = Pt(11)
            
            # Add sections or full content
            doc.add_page_break()
            
            if report_data.get('sections'):
                section_titles = {
                    'executive_summary': '执行摘要',
                    'industry_overview': '产业概览',
                    'policy_landscape': '政策环境',
                    'ecosystem': '产业生态',
                    'value_chain': '产业链分析',
                    'ai_integration': 'AI融合潜力',
                    'conclusion': '结论与建议'
                }
                
                for section_key, section_content in report_data['sections'].items():
                    doc.add_heading(section_titles.get(section_key, section_key), level=1)
                    for para in section_content.split('\n\n'):
                        if para.strip():
                            # Remove markdown and JSON formatting
                            clean_para = self._clean_content_formatting(para.strip())
                            p = doc.add_paragraph(clean_para)
                            p.style.font.size = Pt(11)
                    doc.add_paragraph()  # Empty line
            elif report_data.get('full_content'):
                doc.add_heading('完整报告', level=1)
                full_content = report_data.get('full_content', '')
                for para in full_content.split('\n\n'):
                    if para.strip() and not para.startswith('#'):
                        clean_para = self._clean_content_formatting(para.strip().replace('#', '').strip())
                        if len(clean_para) > 10:
                            p = doc.add_paragraph(clean_para)
                            p.style.font.size = Pt(11)
            else:
                # Handle uploaded analysis reports (dashboard data)
                doc.add_heading('分析报告', level=1)
                
                # Add summary
                if report_data.get('summary'):
                    doc.add_heading('报告摘要', level=2)
                    summary = report_data['summary']
                    if isinstance(summary, dict):
                        # Add key metrics
                        metrics = []
                        if summary.get('word_count'):
                            metrics.append(f"总字数: {summary['word_count']}")
                        if summary.get('reading_time'):
                            metrics.append(f"阅读时间: {summary['reading_time']} 分钟")
                        if summary.get('categories_analyzed'):
                            metrics.append(f"分析类别: {summary['categories_analyzed']}")
                        
                        if metrics:
                            metrics_text = " | ".join(metrics)
                            metrics_para = doc.add_paragraph(metrics_text)
                            metrics_para.style.font.size = Pt(11)
                            doc.add_paragraph()  # Empty line
                        
                        # Add key highlights
                        if summary.get('key_highlights'):
                            doc.add_heading('关键发现', level=3)
                            for highlight in summary['key_highlights']:
                                p = doc.add_paragraph(highlight, style='List Bullet')
                                p.style.font.size = Pt(11)
                
                # Add categories analysis with scores
                if report_data.get('categories'):
                    doc.add_heading('内容分类分析', level=2)
                    for category_name, category_data in report_data['categories'].items():
                        if category_data.get('has_content'):
                            # Add category name and score
                            score_text = f"{category_name} (得分: {category_data.get('relevance_score', 0)})"
                            doc.add_heading(score_text, level=3)
                            
                            # Add key points
                            if category_data.get('key_points'):
                                for point in category_data['key_points']:
                                    p = doc.add_paragraph(point, style='List Bullet')
                                    p.style.font.size = Pt(11)
                
                # Add AI opportunities with table
                if report_data.get('ai_opportunities'):
                    doc.add_heading('AI应用机会', level=2)
                    
                    # Create table for AI opportunities
                    from docx.shared import Inches
                    table = doc.add_table(rows=1, cols=4)
                    table.style = 'Table Grid'
                    
                    # Set column headers
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = '应用领域'
                    hdr_cells[1].text = '潜力得分'
                    hdr_cells[2].text = '优先级'
                    hdr_cells[3].text = '建议'
                    
                    # Add data rows
                    for ai_name, ai_data_item in report_data['ai_opportunities'].items():
                        score = ai_data_item.get('potential_score', 0)
                        priority = ai_data_item.get('priority_level', 'none')
                        priority_text = {'high': '高', 'medium': '中', 'low': '低', 'none': '无'}.get(priority, priority)
                        recommendation = ai_data_item.get('recommendation', 'N/A')
                        
                        row_cells = table.add_row().cells
                        row_cells[0].text = ai_name
                        row_cells[1].text = str(score)
                        row_cells[2].text = priority_text
                        row_cells[3].text = recommendation
                    
                    doc.add_paragraph()  # Empty line
                
                # Add key insights
                if report_data.get('key_insights'):
                    doc.add_heading('关键洞察', level=2)
                    for insight in report_data['key_insights']:
                        if isinstance(insight, dict):
                            insight_title = insight.get('title', '洞察')
                            doc.add_heading(insight_title, level=3)
                            insight_text = insight.get('text') or str(insight.get('data', ''))
                            if insight_text:
                                p = doc.add_paragraph(insight_text)
                                p.style.font.size = Pt(11)
                        else:
                            p = doc.add_paragraph(str(insight))
                            p.style.font.size = Pt(11)
                        doc.add_paragraph()  # Empty line
                
                # Add charts data as tables
                if report_data.get('charts'):
                    doc.add_page_break()
                    doc.add_heading('数据分析图表', level=1)
                    
                    charts = report_data['charts']
                    
                    # Category distribution chart data
                    if charts.get('category_distribution'):
                        doc.add_heading('内容分布分析', level=2)
                        cat_data = charts['category_distribution']
                        if cat_data.get('data'):
                            table = doc.add_table(rows=1, cols=2)
                            table.style = 'Table Grid'
                            
                            # Set headers
                            hdr_cells = table.rows[0].cells
                            hdr_cells[0].text = '类别'
                            hdr_cells[1].text = '得分'
                            
                            # Add data
                            for i, label in enumerate(cat_data['data'][0].get('labels', [])):
                                value = cat_data['data'][0].get('chart_values', [])[i] if i < len(cat_data['data'][0].get('chart_values', [])) else 0
                                
                                row_cells = table.add_row().cells
                                row_cells[0].text = label
                                row_cells[1].text = str(value)
                    
                    # AI opportunities chart data
                    if charts.get('ai_opportunities'):
                        doc.add_heading('AI应用潜力评估', level=2)
                        ai_chart_data = charts['ai_opportunities']
                        if ai_chart_data.get('data'):
                            table = doc.add_table(rows=1, cols=2)
                            table.style = 'Table Grid'
                            
                            # Set headers
                            hdr_cells = table.rows[0].cells
                            hdr_cells[0].text = '应用领域'
                            hdr_cells[1].text = '潜力得分'
                            
                            # Add data
                            for i, label in enumerate(ai_chart_data['data'][0].get('theta', [])):
                                value = ai_chart_data['data'][0].get('r', [])[i] if i < len(ai_chart_data['data'][0].get('r', [])) else 0
                                
                                row_cells = table.add_row().cells
                                row_cells[0].text = label
                                row_cells[1].text = str(value)
                    
                    # Keyword frequency chart data
                    if charts.get('keyword_frequency'):
                        doc.add_heading('关键词频率分析', level=2)
                        keyword_data = charts['keyword_frequency']
                        if keyword_data.get('data'):
                            table = doc.add_table(rows=1, cols=2)
                            table.style = 'Table Grid'
                            
                            # Set headers
                            hdr_cells = table.rows[0].cells
                            hdr_cells[0].text = '关键词'
                            hdr_cells[1].text = '频率'
                            
                            # Add data (limit to top 15)
                            for i, keyword in enumerate(keyword_data['data'][0].get('x', [])[:15]):
                                freq = keyword_data['data'][0].get('y', [])[i] if i < len(keyword_data['data'][0].get('y', [])) else 0
                                
                                row_cells = table.add_row().cells
                                row_cells[0].text = keyword
                                row_cells[1].text = str(freq)
            
            # Save document
            doc.save(str(output_path))
            
            logger.info(f"Word document exported successfully: {output_path}")
            return str(output_path)
        
        except Exception as e:
            logger.error(f"Error exporting to Word: {e}")
            raise
    
    def export_to_excel(self, report_data: Dict, filename: str) -> str:
        """Export report data to Excel format.
        
        Args:
            report_data: Report data dictionary
            filename: Output filename (without extension)
        
        Returns:
            Path to the generated Excel file
        """
        try:
            output_path = self.output_dir / f"{filename}.xlsx"
            
            # Create Excel workbook
            wb = Workbook()
            
            # Summary sheet
            ws_summary = wb.active
            ws_summary.title = "摘要"
            
            # Add header
            ws_summary['A1'] = f"{report_data.get('city', 'N/A')} {report_data.get('industry', 'N/A')} 产业分析报告"
            ws_summary['A1'].font = Font(size=16, bold=True, color="667eea")
            ws_summary['A2'] = f"生成时间: {report_data.get('generated_at', 'N/A')}"
            
            # Add summary
            if report_data.get('summary'):
                ws_summary['A4'] = "执行摘要"
                ws_summary['A4'].font = Font(size=14, bold=True)
                ws_summary['A5'] = report_data['summary'].get('zh', '')
                ws_summary.column_dimensions['A'].width = 80
            
            # SWOT sheet
            if report_data.get('swot_analysis'):
                ws_swot = wb.create_sheet("SWOT分析")
                swot = report_data['swot_analysis']
                
                # Headers
                headers = ['类型', '内容']
                ws_swot.append(headers)
                
                # Data
                for strength in swot.get('strengths', []):
                    ws_swot.append(['优势', strength])
                for weakness in swot.get('weaknesses', []):
                    ws_swot.append(['劣势', weakness])
                for opportunity in swot.get('opportunities', []):
                    ws_swot.append(['机遇', opportunity])
                for threat in swot.get('threats', []):
                    ws_swot.append(['威胁', threat])
                
                # Format headers
                for cell in ws_swot[1]:
                    cell.font = Font(bold=True, color="FFFFFF")
                    cell.fill = PatternFill(start_color="667eea", end_color="667eea", fill_type="solid")
                
                ws_swot.column_dimensions['A'].width = 15
                ws_swot.column_dimensions['B'].width = 60
            
            # Save workbook
            wb.save(str(output_path))
            
            logger.info(f"Excel file exported successfully: {output_path}")
            return str(output_path)
        
        except Exception as e:
            logger.error(f"Error exporting to Excel: {e}")
            raise
